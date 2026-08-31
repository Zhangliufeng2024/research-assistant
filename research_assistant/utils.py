"""Utility functions for research assistant."""

import re
from pathlib import Path
from typing import Any


def find_existing_papers(output_folder: Path) -> list[dict[str, Any]]:
    """
    Get all existing paper directories with their metadata.

    Args:
        output_folder: Path to the paper outputs folder.

    Returns:
        List of dicts with path, name, and timestamp info.
    """
    papers = []
    if not output_folder.exists():
        return papers

    for paper_dir in output_folder.iterdir():
        if paper_dir.is_dir():
            papers.append({
                'path': paper_dir,
                'name': paper_dir.name,
                'mtime': paper_dir.stat().st_mtime
            })

    # Sort by modification time (most recent first)
    papers.sort(key=lambda x: x['mtime'], reverse=True)
    return papers


def detect_paper_reference(user_input: str, existing_papers: list[dict[str, Any]]) -> Path | None:
    """
    Try to detect if the user is referring to an existing paper.

    Args:
        user_input: User's input text.
        existing_papers: List of existing paper dictionaries.

    Returns:
        The paper path if found, None otherwise.
    """
    if not existing_papers:
        return None

    user_input_lower = user_input.lower()

    # Keywords that suggest continuing with existing work
    continuation_keywords = [
        "continue", "update", "edit", "revise", "modify", "change",
        "add to", "fix", "improve", "review", "the paper", "this paper",
        "my paper", "current paper", "previous paper", "last paper",
        "poster", "the poster", "my poster", "presentation", "the presentation",
        "my presentation", "previous presentation", "last presentation",
        "compile", "generate pdf"
    ]

    # Keywords that suggest searching for/looking up an existing paper
    search_keywords = [
        "look for", "find", "search for", "where is", "which paper",
        "show me", "open", "locate", "get"
    ]

    # Keywords that explicitly indicate a new paper
    new_paper_keywords = [
        "new paper", "start fresh", "start afresh", "create new",
        "different paper", "another paper", "write a new",
        "new presentation", "new poster", "different presentation", "another presentation"
    ]

    # If user explicitly wants a new paper, return None
    if any(keyword in user_input_lower for keyword in new_paper_keywords):
        return None

    # Check if user mentions continuation or search keywords
    has_continuation_keyword = any(keyword in user_input_lower for keyword in continuation_keywords)
    has_search_keyword = any(keyword in user_input_lower for keyword in search_keywords)

    # Try to find paper by name/topic keywords
    best_match = None
    best_match_score = 0

    for paper in existing_papers:
        paper_name = paper['name'].lower()
        # Extract topic from directory name (format: YYYYMMDD_HHMMSS_topic)
        parts = paper_name.split('_', 2)
        if len(parts) >= 3:
            topic = parts[2].replace('_', ' ')
            # Check if topic words appear in user input
            topic_words = topic.split()
            matches = sum(1 for word in topic_words if len(word) > 3 and word in user_input_lower)

            # Keep track of best match
            if matches > best_match_score:
                best_match_score = matches
                best_match = paper['path']

            # If we have a strong match (2+ topic words), return it
            # This is especially important for search keywords
            if matches >= 2 and (has_search_keyword or has_continuation_keyword):
                return paper['path']

    # If we found any match with search keywords, return the best one
    if has_search_keyword and best_match_score > 0:
        return best_match

    # If user used continuation keywords but no specific match, use most recent paper
    if has_continuation_keyword and existing_papers:
        return existing_papers[0]['path']

    return None


def scan_paper_directory(paper_dir: Path) -> dict[str, Any]:
    """
    Scan a paper directory and collect all file information.

    Args:
        paper_dir: Path to the paper directory.

    Returns:
        Dictionary with comprehensive file information.
    """
    result = {
        'pdf_final': None,
        'tex_final': None,
        'pdf_drafts': [],
        'tex_drafts': [],
        'docx_final': None,
        'docx_drafts': [],
        'bibliography': None,
        'figures': [],
        'data': [],
        'sources': [],
        'progress_log': None,
        'summary': None,
    }

    if not paper_dir.exists():
        return result

    # Scan final/ directory
    final_dir = paper_dir / "final"
    if final_dir.exists():
        for file in final_dir.iterdir():
            if file.is_file():
                if file.suffix == '.pdf':
                    result['pdf_final'] = str(file)
                elif file.suffix == '.tex':
                    result['tex_final'] = str(file)
                elif file.suffix == '.docx':
                    result['docx_final'] = str(file)

    # Scan drafts/ directory
    drafts_dir = paper_dir / "drafts"
    if drafts_dir.exists():
        for file in sorted(drafts_dir.iterdir()):
            if file.is_file():
                if file.suffix == '.pdf':
                    result['pdf_drafts'].append(str(file))
                elif file.suffix == '.tex':
                    result['tex_drafts'].append(str(file))
                elif file.suffix == '.docx':
                    result['docx_drafts'].append(str(file))

    # Scan references/ directory
    references_dir = paper_dir / "references"
    if references_dir.exists():
        bib_file = references_dir / "references.bib"
        if bib_file.exists():
            result['bibliography'] = str(bib_file)

    # Scan figures/ directory
    figures_dir = paper_dir / "figures"
    if figures_dir.exists():
        for file in sorted(figures_dir.iterdir()):
            if file.is_file():
                result['figures'].append(str(file))

    # Scan data/ directory
    data_dir = paper_dir / "data"
    if data_dir.exists():
        for file in sorted(data_dir.iterdir()):
            if file.is_file():
                result['data'].append(str(file))

    # Scan sources/ directory
    sources_dir = paper_dir / "sources"
    if sources_dir.exists():
        for file in sorted(sources_dir.iterdir()):
            if file.is_file():
                result['sources'].append(str(file))

    # Check for progress.md and SUMMARY.md
    progress_file = paper_dir / "progress.md"
    if progress_file.exists():
        result['progress_log'] = str(progress_file)

    summary_file = paper_dir / "SUMMARY.md"
    if summary_file.exists():
        result['summary'] = str(summary_file)

    return result


def count_citations_in_bib(bib_file: str | None) -> int:
    """
    Count the number of citations in a BibTeX file.

    Args:
        bib_file: Path to the .bib file.

    Returns:
        Number of citations found.
    """
    if not bib_file or not Path(bib_file).exists():
        return 0

    try:
        with open(bib_file, encoding='utf-8') as f:
            content = f.read()
            # Count @article, @book, @inproceedings, etc.
            # Only match lines that actually start with @ (ignore commented-out entries).
            count = 0
            for line in content.splitlines():
                stripped = line.strip()
                if stripped.startswith('@') and not stripped.startswith('@comment'):
                    count += 1
            return count
    except Exception:
        return 0


def extract_citation_style(bib_file: str | None) -> str:
    """返回固定的引用风格常量（**当前是占位实现**）。

    ⚠️ 行为如实声明（P2-10 修正：旧 docstring 声称会「尝试提取」，实际从不
    解析）：本函数无条件返回 ``"BibTeX"``，**不读取** *bib_file* 的任何内容。
    因此 ``api.generate_paper`` 组装结果里展示给前端的「引用风格」目前是
    常量占位，不是从文献检测出来的真实风格。

    若要实现真提取（解析 \\bibliographystyle 命令 / .cls 文件名），属于功能
    新增，需另行排期——修 docstring 不改变行为是刻意的，避免把功能开发混进
    修订批次。
    """
    return "BibTeX"


def count_words_in_tex(tex_file: str | None) -> int | None:
    """
    Estimate word count in a LaTeX file.

    Args:
        tex_file: Path to the .tex file.

    Returns:
        Estimated word count, or None if file doesn't exist.
    """
    if not tex_file or not Path(tex_file).exists():
        return None

    try:
        with open(tex_file, encoding='utf-8') as f:
            content = f.read()

            # Remove comments first
            content = re.sub(r'%.*', '', content)
            # Iteratively remove LaTeX commands with arguments until stable
            # (handles nesting like \textbf{\textit{word}})
            prev = None
            while prev != content:
                prev = content
                content = re.sub(r'\\[a-zA-Z]+(\*)?(\[.*?\])?(\{[^{}]*\})?', '', content)
            # Remove remaining braces and special chars
            content = re.sub(r'[{}$\\]', '', content)

            words = content.split()
            return len(words)
    except Exception:
        return None


def extract_title_from_tex(tex_file: str | None) -> str | None:
    """
    Extract title from a LaTeX file.

    Args:
        tex_file: Path to the .tex file.

    Returns:
        Title string, or None if not found.
    """
    if not tex_file or not Path(tex_file).exists():
        return None

    try:
        with open(tex_file, encoding='utf-8') as f:
            content = f.read()

            # Look for \title{...} — use DOTALL to handle multi-line titles
            match = re.search(r'\\title\s*\{([^}]+)\}', content, re.DOTALL)
            if match:
                title = match.group(1)
                # Clean up LaTeX commands in title
                title = re.sub(r'\\[a-zA-Z]+(\[.*?\])?(\{.*?\})?', '', title)
                return title.strip()
    except Exception:  # noqa: BLE001 — 尽力而为：标题提取失败返回 None，调用方自行降级
        pass

    return None


def count_words_in_docx(docx_file: str | None) -> int | None:
    """Estimate word count in a .docx file.

    Args:
        docx_file: Path to the .docx file.

    Returns:
        Estimated word count, or None if file doesn't exist.
    """
    if not docx_file or not Path(docx_file).exists():
        return None

    try:
        from docx import Document
        doc = Document(docx_file)
        text = " ".join(p.text for p in doc.paragraphs)
        return len(text.split())
    except Exception:
        return None


def extract_title_from_docx(docx_file: str | None) -> str | None:
    """Extract title from a .docx file.

    Args:
        docx_file: Path to the .docx file.

    Returns:
        Title string, or None if not found.
    """
    if not docx_file or not Path(docx_file).exists():
        return None

    try:
        from docx import Document
        doc = Document(docx_file)
        for p in doc.paragraphs:
            if p.style and p.style.name == "Title":
                if p.text.strip():
                    return p.text.strip()
        for p in doc.paragraphs:
            if p.text.strip():
                return p.text.strip()
    except Exception:  # noqa: BLE001 — 尽力而为：docx 解析失败返回 None，不阻断调用方
        pass

    return None

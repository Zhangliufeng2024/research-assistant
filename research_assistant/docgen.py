"""Document generation helpers for creating Word (.docx) documents.

Provides PaperBuilder — a high-level API for constructing scientific papers
with professional formatting, figure embedding, tables, numbered references,
cross-referencing, and venue-specific templates.
"""

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Document template system
# ---------------------------------------------------------------------------

@dataclass
class DocumentTemplate:
    """Defines page layout, fonts, and formatting for a venue/style."""

    name: str
    display_name: str

    page_width: float = 8.5
    page_height: float = 11.0
    margin_top: float = 1.0
    margin_bottom: float = 1.0
    margin_left: float = 1.0
    margin_right: float = 1.0

    body_font: str = "Times New Roman"
    body_font_size: float = 11.0
    line_spacing: float = 1.15
    space_after_paragraph: float = 6.0

    heading_font: str = "Times New Roman"
    heading1_size: float = 14.0
    heading2_size: float = 12.0
    heading3_size: float = 11.0
    heading_color: tuple[int, int, int] = (0, 0, 0)
    heading_numbered: bool = False

    title_font_size: float = 16.0
    author_font_size: float = 12.0

    abstract_font_size: float = 10.0
    abstract_indent: float = 0.5
    abstract_italic: bool = True

    caption_font_size: float = 10.0
    reference_font_size: float = 10.0

    columns: int = 1

    figure_width: float = 5.5


TEMPLATES: dict[str, DocumentTemplate] = {
    "default": DocumentTemplate(
        name="default",
        display_name="Default (Letter, single column)",
    ),
    "ieee_journal": DocumentTemplate(
        name="ieee_journal",
        display_name="IEEE Transactions Journal",
        margin_top=0.75,
        margin_bottom=0.75,
        margin_left=0.625,
        margin_right=0.625,
        body_font_size=10.0,
        line_spacing=1.0,
        space_after_paragraph=4.0,
        heading1_size=12.0,
        heading2_size=10.0,
        heading3_size=10.0,
        heading_numbered=True,
        title_font_size=24.0,
        author_font_size=11.0,
        abstract_font_size=9.0,
        abstract_indent=0.0,
        abstract_italic=True,
        caption_font_size=8.0,
        reference_font_size=8.0,
        columns=2,
        figure_width=3.25,
    ),
    "ieee_conference": DocumentTemplate(
        name="ieee_conference",
        display_name="IEEE Conference",
        margin_top=0.75,
        margin_bottom=1.0,
        margin_left=0.625,
        margin_right=0.625,
        body_font_size=10.0,
        line_spacing=1.0,
        space_after_paragraph=4.0,
        heading1_size=12.0,
        heading2_size=10.0,
        heading3_size=10.0,
        heading_numbered=True,
        title_font_size=24.0,
        author_font_size=11.0,
        abstract_font_size=9.0,
        abstract_indent=0.0,
        abstract_italic=True,
        caption_font_size=8.0,
        reference_font_size=8.0,
        columns=2,
        figure_width=3.25,
    ),
    "nature": DocumentTemplate(
        name="nature",
        display_name="Nature Article",
        body_font="Arial",
        body_font_size=11.0,
        line_spacing=1.5,
        space_after_paragraph=6.0,
        heading_font="Arial",
        heading1_size=14.0,
        heading2_size=12.0,
        heading3_size=11.0,
        heading_numbered=False,
        title_font_size=18.0,
        author_font_size=12.0,
        abstract_font_size=10.0,
        abstract_indent=0.0,
        abstract_italic=False,
        caption_font_size=10.0,
        reference_font_size=9.0,
        columns=1,
        figure_width=5.5,
    ),
    "apa7": DocumentTemplate(
        name="apa7",
        display_name="APA 7th Edition",
        margin_top=1.0,
        margin_bottom=1.0,
        margin_left=1.0,
        margin_right=1.0,
        body_font="Times New Roman",
        body_font_size=12.0,
        line_spacing=2.0,
        space_after_paragraph=0.0,
        heading_font="Times New Roman",
        heading1_size=12.0,
        heading2_size=12.0,
        heading3_size=12.0,
        heading_numbered=False,
        title_font_size=12.0,
        author_font_size=12.0,
        abstract_font_size=12.0,
        abstract_indent=0.5,
        abstract_italic=False,
        caption_font_size=12.0,
        reference_font_size=12.0,
        columns=1,
        figure_width=5.5,
    ),
    "asce": DocumentTemplate(
        name="asce",
        display_name="ASCE Journal",
        margin_top=1.0,
        margin_bottom=1.0,
        margin_left=1.0,
        margin_right=1.0,
        body_font="Times New Roman",
        body_font_size=12.0,
        line_spacing=2.0,
        space_after_paragraph=0.0,
        heading_font="Times New Roman",
        heading1_size=12.0,
        heading2_size=12.0,
        heading3_size=12.0,
        heading_numbered=False,
        title_font_size=14.0,
        author_font_size=12.0,
        abstract_font_size=12.0,
        abstract_indent=0.5,
        abstract_italic=False,
        caption_font_size=10.0,
        reference_font_size=10.0,
        columns=1,
        figure_width=5.5,
    ),
    "elsevier": DocumentTemplate(
        name="elsevier",
        display_name="Elsevier Journal",
        margin_top=1.0,
        margin_bottom=1.0,
        margin_left=1.0,
        margin_right=1.0,
        body_font="Times New Roman",
        body_font_size=11.0,
        line_spacing=1.5,
        space_after_paragraph=6.0,
        heading_font="Times New Roman",
        heading1_size=14.0,
        heading2_size=12.0,
        heading3_size=11.0,
        heading_numbered=True,
        title_font_size=17.0,
        author_font_size=12.0,
        abstract_font_size=10.0,
        abstract_indent=0.0,
        abstract_italic=False,
        caption_font_size=9.0,
        reference_font_size=9.0,
        columns=1,
        figure_width=5.5,
    ),
}


def get_template(name: str) -> DocumentTemplate:
    """Return a template by name.  Falls back to 'default' for unknown names."""
    return TEMPLATES.get(name, TEMPLATES["default"])


def list_templates() -> list[str]:
    """Return all available template names."""
    return list(TEMPLATES.keys())


# ---------------------------------------------------------------------------
# Cross-reference manager
# ---------------------------------------------------------------------------

_REF_PATTERN = re.compile(r"\{ref:([^}]+)\}")


@dataclass
class _RefEntry:
    ref_type: str   # "Figure", "Table", "Section"
    number: int     # 1, 2, 3, ...
    title: str = "" # optional display text


class CrossReferenceManager:
    """Tracks labels for figures, tables, and sections and resolves {ref:label} markers."""

    def __init__(self) -> None:
        self._entries: dict[str, _RefEntry] = {}

    def register(self, label: str, ref_type: str, number: int, title: str = "") -> None:
        self._entries[label] = _RefEntry(ref_type=ref_type, number=number, title=title)

    def get(self, label: str) -> Optional[_RefEntry]:
        return self._entries.get(label)

    def format_ref(self, label: str) -> str:
        entry = self._entries.get(label)
        if entry is None:
            return f"[??{label}]"
        return f"{entry.ref_type} {entry.number}"

    def resolve(self, text: str) -> str:
        """Replace all {ref:label} markers in *text* with formatted references."""
        def _repl(m: re.Match) -> str:
            return self.format_ref(m.group(1))
        return _REF_PATTERN.sub(_repl, text)


# ---------------------------------------------------------------------------
# Reference data model
# ---------------------------------------------------------------------------

@dataclass
class Reference:
    """A single bibliographic reference."""
    key: str
    authors: str
    title: str
    year: int
    journal: str = ""
    volume: str = ""
    pages: str = ""
    doi: str = ""
    url: str = ""
    booktitle: str = ""
    publisher: str = ""
    ref_type: str = "article"

    def format_numbered(self, number: int) -> str:
        """Format as a numbered reference string (IEEE-like)."""
        parts = [f"[{number}]"]
        if self.authors:
            parts.append(f"{self.authors}.")
        parts.append(f'"{self.title}."')
        if self.journal:
            parts.append(f"{self.journal},")
        elif self.booktitle:
            parts.append(f"in {self.booktitle},")
        if self.volume:
            vol = f"vol. {self.volume}"
            if self.pages:
                vol += f", pp. {self.pages}"
            parts.append(f"{vol},")
        elif self.pages:
            parts.append(f"pp. {self.pages},")
        parts.append(f"{self.year}.")
        if self.doi:
            parts.append(f"DOI: {self.doi}")
        return " ".join(parts)

    # keep legacy alias
    format_apa = format_numbered


# ---------------------------------------------------------------------------
# Helper
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, color: str) -> None:
    """Set background shading on a table cell."""
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): color,
    })
    shading.append(shd)


# ---------------------------------------------------------------------------
# PaperBuilder
# ---------------------------------------------------------------------------

class PaperBuilder:
    """High-level builder for scientific Word documents.

    Usage::

        paper = PaperBuilder("My Title", authors=["Alice", "Bob"])
        paper.add_section("Introduction", "Body text here...")
        paper.add_figure("figures/fig1.png", "System overview.", label="fig:overview")
        paper.add_table(["A", "B"], [["1", "2"]], caption="Results", label="tab:results")
        paper.add_section("Discussion", "As shown in {ref:fig:overview} and {ref:tab:results}...")
        paper.add_references([Reference(key="smith2023", ...)])
        paper.save("final/manuscript.docx")

    Templates::

        paper = PaperBuilder("Title", template="ieee_journal")
        print(PaperBuilder.available_templates())

    Cross-references are resolved automatically during ``save()``.
    Use ``{ref:label}`` in content text to insert "Figure 1", "Table 2", etc.
    """

    def __init__(
        self,
        title: str,
        authors: Optional[list[str]] = None,
        abstract: Optional[str] = None,
        affiliation: Optional[str] = None,
        template: str | DocumentTemplate = "default",
    ):
        if isinstance(template, str):
            self._template = get_template(template)
        else:
            self._template = template

        self.doc = Document()
        self._figure_count = 0
        self._table_count = 0
        self._section_counters: list[int] = [0, 0, 0]  # level 1, 2, 3
        self._refs = CrossReferenceManager()
        self._setup_styles()
        self._add_title_block(title, authors, affiliation)
        if abstract:
            self._add_abstract(abstract)

    # -- public helpers --

    @staticmethod
    def available_templates() -> list[str]:
        """Return all registered template names."""
        return list_templates()

    @property
    def template(self) -> DocumentTemplate:
        return self._template

    @property
    def refs(self) -> CrossReferenceManager:
        """Access the cross-reference manager for manual lookups."""
        return self._refs

    def ref(self, label: str) -> str:
        """Return the formatted reference string for *label* (e.g. ``"Figure 1"``)."""
        return self._refs.format_ref(label)

    # -- styles --

    def _setup_styles(self) -> None:
        """Configure page layout and text styles from the active template."""
        t = self._template

        for section in self.doc.sections:
            section.page_width = Inches(t.page_width)
            section.page_height = Inches(t.page_height)
            section.top_margin = Inches(t.margin_top)
            section.bottom_margin = Inches(t.margin_bottom)
            section.left_margin = Inches(t.margin_left)
            section.right_margin = Inches(t.margin_right)

            if t.columns > 1:
                sect_pr = section._sectPr
                cols_elem = sect_pr.makeelement(qn("w:cols"), {
                    qn("w:num"): str(t.columns),
                    qn("w:space"): "720",
                })
                sect_pr.append(cols_elem)

        style = self.doc.styles["Normal"]
        style.font.size = Pt(t.body_font_size)
        style.font.name = t.body_font
        style.paragraph_format.line_spacing = t.line_spacing
        style.paragraph_format.space_after = Pt(t.space_after_paragraph)

        for level in range(1, 4):
            name = f"Heading {level}"
            if name in self.doc.styles:
                h = self.doc.styles[name]
                h.font.name = t.heading_font
                h.font.color.rgb = RGBColor(*t.heading_color)
                if level == 1:
                    h.font.size = Pt(t.heading1_size)
                    h.font.bold = True
                elif level == 2:
                    h.font.size = Pt(t.heading2_size)
                    h.font.bold = True
                else:
                    h.font.size = Pt(t.heading3_size)
                    h.font.bold = True
                    h.font.italic = True

    # -- title block --

    def _add_title_block(
        self,
        title: str,
        authors: Optional[list[str]],
        affiliation: Optional[str],
    ) -> None:
        t = self._template

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(t.title_font_size)
        run.font.name = t.heading_font
        p.paragraph_format.space_after = Pt(4)

        if authors:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(", ".join(authors))
            run.font.size = Pt(t.author_font_size)
            run.font.name = t.body_font
            p.paragraph_format.space_after = Pt(2)

        if affiliation:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(affiliation)
            run.font.size = Pt(10)
            run.font.italic = True
            run.font.name = t.body_font
            p.paragraph_format.space_after = Pt(12)

    # -- abstract --

    def _add_abstract(self, text: str) -> None:
        t = self._template

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("Abstract")
        run.bold = True
        run.font.size = Pt(t.body_font_size)
        run.font.name = t.heading_font
        p.paragraph_format.space_after = Pt(4)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        run.italic = t.abstract_italic
        run.font.size = Pt(t.abstract_font_size)
        run.font.name = t.body_font
        pf = p.paragraph_format
        if t.abstract_indent > 0:
            pf.left_indent = Inches(t.abstract_indent)
            pf.right_indent = Inches(t.abstract_indent)
        pf.space_after = Pt(12)

    # -- sections --

    def add_section(
        self,
        title: str,
        content: str,
        level: int = 1,
        label: Optional[str] = None,
    ) -> None:
        """Add a section with heading and body text.

        Args:
            title: Section heading text.
            content: Body text (plain text; double-newlines become new paragraphs).
            level: Heading level (1, 2, or 3).
            label: Optional label for cross-referencing (e.g. ``"sec:intro"``).
        """
        level = max(1, min(3, level))
        t = self._template

        # Update counters for numbered headings
        self._section_counters[level - 1] += 1
        for deeper in range(level, 3):
            self._section_counters[deeper] = 0

        sec_number = ".".join(
            str(self._section_counters[i]) for i in range(level)
        )

        if label:
            self._refs.register(label, "Section", self._section_counters[0], title=title)

        heading_text = f"{sec_number}. {title}" if t.heading_numbered else title
        self.doc.add_heading(heading_text, level=level)

        for para_text in content.split("\n\n"):
            para_text = para_text.strip()
            if not para_text:
                continue
            p = self.doc.add_paragraph(para_text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # -- figures --

    def add_figure(
        self,
        image_path: str,
        caption: str,
        label: Optional[str] = None,
        width_inches: Optional[float] = None,
    ) -> None:
        """Add a figure with auto-numbered caption.

        Args:
            image_path: Path to the image file.
            caption: Caption text.
            label: Optional label for cross-referencing (e.g. ``"fig:overview"``).
            width_inches: Image width.  Defaults to the template's ``figure_width``.
        """
        path = Path(image_path)
        if not path.exists():
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Missing figure: {image_path}]")
            run.italic = True
            return

        self._figure_count += 1
        t = self._template
        width = width_inches if width_inches is not None else t.figure_width

        if label:
            self._refs.register(label, "Figure", self._figure_count)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))

        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(f"Figure {self._figure_count}. {caption}")
        run.italic = True
        run.font.size = Pt(t.caption_font_size)
        run.font.name = t.body_font
        cap.paragraph_format.space_after = Pt(12)

    # -- tables --

    def add_table(
        self,
        headers: list[str],
        rows: list[list[str]],
        caption: Optional[str] = None,
        label: Optional[str] = None,
    ) -> None:
        """Add a table with optional auto-numbered caption.

        Args:
            headers: Column header strings.
            rows: List of rows, each a list of cell strings.
            caption: Optional caption text.
            label: Optional label for cross-referencing (e.g. ``"tab:results"``).
        """
        self._table_count += 1
        t = self._template

        if label:
            self._refs.register(label, "Table", self._table_count)

        if caption:
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(12)
            run = cap.add_run(f"Table {self._table_count}. {caption}")
            run.bold = True
            run.font.size = Pt(t.caption_font_size)
            run.font.name = t.body_font

        n_cols = len(headers)
        table = self.doc.add_table(rows=1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(t.caption_font_size)
                    run.font.name = t.body_font
            _set_cell_shading(cell, "D9E2F3")

        for row_data in rows:
            row = table.add_row()
            for i, cell_text in enumerate(row_data):
                if i < n_cols:
                    cell = row.cells[i]
                    cell.text = str(cell_text)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(t.caption_font_size)
                            run.font.name = t.body_font

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)

    # -- references --

    def add_references(self, references: list[Reference]) -> None:
        """Add a numbered References section.

        Args:
            references: List of Reference objects to format.
        """
        t = self._template
        self.doc.add_heading("References", level=1)

        for i, ref in enumerate(references, 1):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.first_line_indent = Inches(-0.4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(ref.format_numbered(i))
            run.font.size = Pt(t.reference_font_size)
            run.font.name = t.body_font

    def add_references_from_bibtex(self, bib_path: str) -> None:
        """Parse a .bib file and add all entries as numbered references."""
        refs = parse_bibtex(bib_path)
        if refs:
            self.add_references(refs)

    # -- cross-reference resolution --

    def _resolve_all_references(self) -> None:
        """Scan every paragraph and table cell, resolving {ref:label} markers."""
        for para in self.doc.paragraphs:
            self._resolve_paragraph(para)
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._resolve_paragraph(para)

    def _resolve_paragraph(self, para) -> None:
        for run in para.runs:
            if "{ref:" in run.text:
                run.text = self._refs.resolve(run.text)

    # -- save --

    def save(self, path: str) -> str:
        """Save the document.  Cross-references are resolved before writing.

        Returns:
            Absolute path to the saved file.
        """
        self._resolve_all_references()
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(p))
        return str(p.resolve())


# ---------------------------------------------------------------------------
# BibTeX parser (simple regex-based, sufficient for machine-generated entries)
# ---------------------------------------------------------------------------

_ENTRY_RE = re.compile(
    r"@(\w+)\s*\{\s*([^,\s]+)\s*,(.+?)\n\s*\}",
    re.DOTALL,
)

_FIELD_RE = re.compile(
    r"(\w+)\s*=\s*\{([^}]*)\}",
)


def parse_bibtex(bib_path: str) -> list[Reference]:
    """Parse a .bib file into a list of Reference objects."""
    path = Path(bib_path)
    if not path.exists():
        return []

    try:
        text = path.read_text(encoding="utf-8")
    except Exception:
        return []

    refs: list[Reference] = []
    for match in _ENTRY_RE.finditer(text):
        ref_type = match.group(1).lower()
        key = match.group(2)
        body = match.group(3)

        fields: dict[str, str] = {}
        for fm in _FIELD_RE.finditer(body):
            fields[fm.group(1).lower()] = fm.group(2).strip()

        year_str = fields.get("year", "0")
        try:
            year = int(year_str)
        except ValueError:
            year = 0

        refs.append(Reference(
            key=key,
            authors=fields.get("author", ""),
            title=fields.get("title", ""),
            year=year,
            journal=fields.get("journal", ""),
            volume=fields.get("volume", ""),
            pages=fields.get("pages", ""),
            doi=fields.get("doi", ""),
            url=fields.get("url", ""),
            booktitle=fields.get("booktitle", ""),
            publisher=fields.get("publisher", ""),
            ref_type=ref_type,
        ))

    return refs

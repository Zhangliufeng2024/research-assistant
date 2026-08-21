"""Document completeness gate — no placeholders, sections present, figures exist."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Optional

from .base import Gate, GateResult

_PLACEHOLDER_RE = re.compile(
    r"\bTODO\b|PLACEHOLDER|Lorem ipsum|\[citation needed\]|\[INSERT", re.IGNORECASE
)
_FIGURE_REF_RE = re.compile(r"figures/([\w.\-]+\.(?:png|jpg|jpeg|svg))", re.IGNORECASE)


def extract_docx_text(docx_path: str | Path) -> str:
    """Concatenated paragraph text of a .docx file."""
    from docx import Document

    doc = Document(str(docx_path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def count_words(text: str) -> int:
    return len(re.findall(r"\S+", text))


class DocGate(Gate):
    """Check a generated .docx for completeness before finalization.

    Blocking checks:
      - file exists and is readable
      - every expected section title appears in the document text
      - figure paths referenced in the text exist on disk
      - no placeholder markers (TODO / Lorem ipsum / [citation needed] ...)
      - word count >= ratio * target_words (when target_words given)
    """

    name = "document"

    def __init__(
        self,
        docx_path: str | Path,
        expected_sections: Optional[list[str]] = None,
        figures_dir: Optional[str | Path] = None,
        target_words: Optional[int] = None,
        min_words_ratio: float = 0.8,
    ) -> None:
        self.docx_path = Path(docx_path)
        self.expected_sections = expected_sections or []
        self.figures_dir = Path(figures_dir) if figures_dir else None
        self.target_words = target_words
        self.min_words_ratio = min_words_ratio

    async def run(self, context: dict[str, Any]) -> GateResult:
        failures: list[str] = []
        details: dict[str, Any] = {"docx": str(self.docx_path)}

        if not self.docx_path.exists():
            return self._result(False, [f"document not found: {self.docx_path}"])

        try:
            text = extract_docx_text(self.docx_path)
        except Exception as e:
            return self._result(False, [f"cannot read document: {e}"])

        lowered = re.sub(r"\s+", " ", text).lower()

        # 1. Sections present
        missing = [
            s for s in self.expected_sections
            if re.sub(r"\s+", " ", s).lower() not in lowered
        ]
        if missing:
            failures.append(f"missing sections: {', '.join(missing)}")
        details["sections_expected"] = len(self.expected_sections)
        details["sections_missing"] = missing

        # 2. Referenced figures exist
        referenced = _FIGURE_REF_RE.findall(text)
        missing_figs = []
        if self.figures_dir is not None:
            for name in sorted(set(referenced)):
                if not (self.figures_dir / name).exists():
                    missing_figs.append(name)
            if missing_figs:
                failures.append(f"figures referenced but missing on disk: {', '.join(missing_figs)}")
        details["figures_referenced"] = sorted(set(referenced))
        details["figures_missing"] = missing_figs

        # 3. Placeholders
        placeholders = _PLACEHOLDER_RE.findall(text)
        if placeholders:
            failures.append(
                f"placeholder text found ({len(placeholders)}x): "
                f"{sorted(set(p.upper() for p in placeholders))}"
            )
        details["placeholders"] = len(placeholders)

        # 4. Word count
        words = count_words(text)
        details["word_count"] = words
        if self.target_words:
            floor = int(self.min_words_ratio * self.target_words)
            details["target_words"] = self.target_words
            details["word_floor"] = floor
            if words < floor:
                failures.append(
                    f"word count {words} below floor {floor} "
                    f"({self.min_words_ratio:.0%} of target {self.target_words})"
                )

        return self._result(not failures, failures or None, details)

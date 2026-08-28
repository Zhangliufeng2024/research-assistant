"""工作区 API（R5 计划 R1：W1-W4；R8 增运行时切换 POST /workspace/root）。

约定：工作区根保存在 ``app.state.cwd``；未初始化 state 的轻量测试才回退
到 ``Path.cwd()``。所有端点的 ``path`` 参数均相对该根解析，且必须通过
:func:`research_assistant.core.safe_resolve` 围栏校验：越界一律 403、
目标不存在 404。路由本身不带 ``/api`` 前缀，由 app.py 以
``prefix="/api"`` 挂载（与 web/routes.py 同一惯例）。

R8 运行时切换（POST /workspace/root）只更新 app state，不改变进程 CWD；
避免并发请求、后台任务和第三方库的相对路径语义被全局切换污染。
"""

import mimetypes
import os
import platform
import subprocess
from pathlib import Path

from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel

from ..artifacts import ArtifactVersionStore
from ..config import load_project_env, resolve_model
from ..core import ensure_output_folder, safe_resolve, setup_claude_skills

router = APIRouter()

#: W2 目录清单忽略项：版本库 / 字节码缓存 / 内部状态 / 前端依赖目录。
_IGNORE_DIRS = frozenset({".git", "__pycache__", ".ra", "node_modules"})

#: W3 文本预览字节上限：超过即截断并置 truncated=True。
TEXT_PREVIEW_LIMIT = 256 * 1024

#: W3 docx 预览最多抽取的正文段落数（表格单元格文本另行全部拼接后再统一截断）。
DOCX_PARAGRAPH_LIMIT = 200


def _root_of(request: Request) -> Path:
    """Return the request-scoped application workspace root."""
    return Path(getattr(request.app.state, "cwd", None) or Path.cwd()).resolve()


def _version_store(request: Request) -> ArtifactVersionStore:
    return ArtifactVersionStore(_root_of(request))


@router.get("/workspace/changes")
async def list_workspace_changes(request: Request, limit: int = 200):
    return _version_store(request).list(limit)


@router.get("/workspace/changes/{change_id}")
async def get_workspace_change(change_id: str, request: Request):
    try:
        return _version_store(request).diff(change_id)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="变更不存在") from exc


class RestoreChangePayload(BaseModel):
    side: str = "before"


@router.post("/workspace/changes/{change_id}/restore")
async def restore_workspace_change(
    change_id: str, payload: RestoreChangePayload, request: Request,
):
    try:
        return _version_store(request).restore(change_id, payload.side)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="变更不存在") from exc
    except ValueError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc

#: W3 视为文本读取的扩展名（小写、含点）；无扩展名的小文件同样按文本处理。
_TEXT_EXTENSIONS = frozenset(
    {
        ".txt",
        ".md",
        ".markdown",
        ".rst",
        ".py",
        ".js",
        ".mjs",
        ".cjs",
        ".ts",
        ".tsx",
        ".jsx",
        ".json",
        ".csv",
        ".tsv",
        ".bib",
        ".log",
        ".yaml",
        ".yml",
        ".tex",
        ".html",
        ".htm",
        ".css",
        ".scss",
        ".xml",
        ".toml",
        ".ini",
        ".cfg",
        ".conf",
        ".sh",
        ".bat",
        ".ps1",
        ".sql",
    }
)


def _fence(path: str | None, root: Path) -> Path:
    """把请求路径解析进工作区根之内，越界抛 403。

    绝对路径、盘符路径、``..`` 回溯（含 URL 编码形式）在「根 + 参数」
    拼接后都会落到根外，统一交给 safe_resolve 拦截。
    """
    raw = (path or "").strip()
    candidate = root / raw if raw else root
    try:
        return safe_resolve(candidate, root)
    except ValueError as exc:
        raise HTTPException(status_code=403, detail="路径越界，拒绝访问") from exc


#: 预览/下载端点的敏感目标前缀：环境变量文件（.env*）、内部状态库目录
#（.ra*，含 platform/sources sqlite）与版本库目录（.git*）。
_PROTECTED_PREFIXES = (".env", ".ra", ".git")


def _is_protected_target(target: Path, root: Path) -> bool:
    """相对路径任一段命中敏感前缀即拒绝（围栏保证 target 在根内）。"""
    try:
        rel = target.relative_to(root)
    except ValueError:
        return True  # 围栏外的异常情况一律视为受保护
    return any(part.lower().startswith(_PROTECTED_PREFIXES) for part in rel.parts)


def _rel_posix(path: Path, root: Path) -> str:
    """相对工作区根的 POSIX 风格显示路径；不在根内时回退绝对路径。

    两侧都先 resolve：入参可能是未归一化的相对/混合分隔符路径
    （如 app.state.output_folder），而 ``relative_to`` 本身不做解析。
    """
    try:
        rel = path.resolve().relative_to(root.resolve()).as_posix()
    except ValueError:
        return path.as_posix()
    return "" if rel == "." else rel  # 根本身显示为空串（树接口回显约定）


def _read_text_head(path: Path, size: int) -> tuple[str, bool]:
    """读文件头部 ≤TEXT_PREVIEW_LIMIT 字节并按 UTF-8 解码（非法字节替换）。"""
    with open(path, "rb") as fh:
        chunk = fh.read(TEXT_PREVIEW_LIMIT)
    return chunk.decode("utf-8", errors="replace"), size > len(chunk)


def _docx_preview(target: Path, size: int) -> dict:
    """抽取 docx 前 DOCX_PARAGRAPH_LIMIT 个段落 + 表格单元格文本作为文本预览。

    解析失败时降级为附件下载，避免单个损坏文档让预览端点报错。
    """
    try:
        from docx import Document  # 延迟导入：python-docx 导入开销较大

        document = Document(str(target))
        parts = [para.text for para in document.paragraphs[:DOCX_PARAGRAPH_LIMIT]]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
    except Exception:
        return FileResponse(
            path=str(target),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{target.name}"'},
        )
    data = "\n".join(parts).encode("utf-8")
    truncated = len(data) > TEXT_PREVIEW_LIMIT
    content = data[:TEXT_PREVIEW_LIMIT].decode("utf-8", errors="replace")
    return {"kind": "text", "content": content, "truncated": truncated, "size": size}


@router.get("/workspace")
async def get_workspace(request: Request):
    """W1 工作区名片：根目录、名称、输出目录显示路径、是否 git 仓库。"""
    root = _root_of(request)
    out_raw = getattr(request.app.state, "output_folder", None)
    output_folder = _rel_posix(Path(out_raw), root) if out_raw is not None else None
    return {
        "root": str(root),
        "name": root.name,
        "output_folder": output_folder,
        "has_git": (root / ".git").exists(),
    }


class SwitchRootPayload(BaseModel):
    path: str


@router.post("/workspace/root")
async def switch_workspace_root(payload: SwitchRootPayload, request: Request):
    """R8 反馈 #1：运行时切换工作区根（Claude Desktop 式界面内添加工作目录）。

    切换只刷新 ``app.state`` 的 cwd / output_folder / model；随后重建输出
    目录与技能、重载配置（全局 .env + 新工作区覆盖层）。

    语义边界（有意为之）：
    - 生成任务运行中仍 409 拒绝，避免 UI 与任务所属工作区产生歧义。
    - 已打开的会话连接不受影响：其 ToolRegistry 在连接时捕获了旧工作区
      的**绝对路径**，进行中的回合继续写原目录（会话属于原工作区，语义
      正确）；新建会话/重连即用新根。
    """
    raw = payload.path.strip().strip('"').strip("'")
    if not raw:
        raise HTTPException(status_code=422, detail="请提供工作目录路径")
    target = Path(os.path.expandvars(raw)).expanduser()
    if not target.is_absolute():
        raise HTTPException(status_code=422, detail="请提供绝对路径")
    target = target.resolve()
    if not target.is_dir():
        raise HTTPException(status_code=404, detail=f"目录不存在：{target}")

    if _root_of(request) == target:
        return {"ok": True, "root": str(target), "unchanged": True}

    tasks_reg = getattr(request.app.state, "active_tasks", None)
    busy = [
        tid
        for tid, t in (tasks_reg or {}).items()
        if not tid.startswith("chat:") and t.get("status") in ("running", "stopping")
    ]
    if busy:
        raise HTTPException(
            status_code=409, detail="有生成任务正在运行，请等待完成（或先停止）再切换工作目录"
        )

    # 技能随新根重同步：冻结包内 .claude 在 research_assistant/ 下，
    # 开发态在仓库根——两级探测与 desktop.bundle_root 的取法互补。
    package_dir = Path(__file__).parent.parent.absolute()
    if not (package_dir / ".claude").exists():
        package_dir = package_dir.parent
    setup_claude_skills(package_dir, target)
    output_folder = ensure_output_folder(target)  # writing_outputs 就绪
    load_project_env(target)  # 全局 + 新工作区 .env
    request.app.state.cwd = target
    request.app.state.output_folder = output_folder
    request.app.state.model = resolve_model(None)
    # Switch the portable project catalog together with the filesystem root.
    # Existing background tasks keep their captured absolute paths and hub;
    # new work is registered in the newly selected project database.
    old_scheduler = getattr(request.app.state, "scheduler", None)
    if old_scheduler is not None:
        # 切换不再等待旧工作区流水线排空（stop 默认 drain=True 会 gather 全部
        # 活跃任务，可能把本 HTTP 请求挂起数小时）。drain=False 只让轮询退出，
        # 活跃任务持有旧 store/hub 的绝对路径引用自然跑完并照常落库。
        await old_scheduler.stop(drain=False)
    from ..runtime import (
        BackgroundTaskHub,
        DurableScheduler,
        PlatformStore,
        build_scheduler_dispatcher,
    )
    platform_store = PlatformStore(target / ".ra" / "platform.sqlite3")
    platform_store.mark_orphaned_running_tasks()
    request.app.state.platform_store = platform_store
    request.app.state.project = platform_store.ensure_project(target)
    request.app.state.task_hub = BackgroundTaskHub(platform_store)
    from ..context import SourceStore
    request.app.state.source_store = SourceStore(target / ".ra" / "sources.sqlite3")
    request.app.state.scheduler = DurableScheduler(platform_store, janitor_cwd=target)
    dispatchers, _dispatch = build_scheduler_dispatcher(
        store=platform_store, hub=request.app.state.task_hub, cwd=target,
        project=request.app.state.project, source_store=request.app.state.source_store,
        default_model=request.app.state.model,
    )
    for workflow_id, dispatcher in dispatchers.items():
        request.app.state.scheduler.register(workflow_id, dispatcher)
    request.app.state.scheduler.start()
    return {"ok": True, "root": str(target), "output_folder": str(output_folder)}


@router.get("/workspace/tree")
async def get_workspace_tree(request: Request, path: str = "", depth: int = 1):
    """W2 目录树单层懒加载：返回一层子项，深层由前端按需逐层展开。

    忽略 ``.git``/``__pycache__``/``.ra``/``node_modules`` 与一切 ``.`` 开头
    的隐藏项；符号链接等解析后落在根外的条目静默跳过（同 routes.list_runs
    策略）。排序：目录在前、名称升序。``depth`` 为预留参数，当前契约是
    单层懒加载，不做递归。
    """
    root = _root_of(request)
    target = _fence(path, root)
    if not target.is_dir():
        raise HTTPException(status_code=404, detail="目录不存在")

    items: list[dict] = []
    try:
        children = list(target.iterdir())
    except OSError as exc:
        raise HTTPException(status_code=404, detail="目录不可读") from exc

    for child in children:
        name = child.name
        if name in _IGNORE_DIRS or name.startswith("."):
            continue
        try:
            resolved = safe_resolve(child, root)
            if not (resolved.is_dir() or resolved.is_file()):
                continue  # 失效链接等非常规条目
            stat = resolved.stat()
        except (ValueError, OSError):
            continue  # 越界或竞态消失的条目直接跳过
        is_dir = resolved.is_dir()
        items.append(
            {
                "name": name,
                "path": _rel_posix(resolved, root),
                "type": "dir" if is_dir else "file",
                "size": None if is_dir else stat.st_size,
                "mtime": stat.st_mtime,
            }
        )

    items.sort(key=lambda item: (item["type"] != "dir", item["name"].lower()))
    return {"path": _rel_posix(target, root), "items": items}


@router.get("/workspace/file")
async def get_workspace_file(request: Request, path: str):
    """W3 泛化文件预览：按类型分流返回。

    - 敏感目标（``.env*`` / ``.ra*`` / ``.git*`` 路径段）一律 403，且先于
      存在性判断 —— 不能借 404/200 探测密钥文件是否存在；
    - 文本类扩展名 / 无扩展名小文件 → 头部 ≤256KB UTF-8 文本 JSON；
    - ``.docx`` → python-docx 抽段落与表格文本，仍为 text kind；
    - 图片 / PDF → FileResponse inline（浏览器内嵌预览）；SVG 除外：
      可携带脚本，强制 attachment 下载；
    - 其余二进制 → FileResponse attachment 下载。
    """
    root = _root_of(request)
    target = _fence(path, root)
    if _is_protected_target(target, root):
        raise HTTPException(status_code=403, detail="受保护文件，不允许预览/下载")
    if not target.is_file():
        raise HTTPException(status_code=404, detail="文件不存在")
    size = target.stat().st_size
    suffix = target.suffix.lower()

    if suffix in _TEXT_EXTENSIONS or suffix == "":
        content, truncated = _read_text_head(target, size)
        return {"kind": "text", "content": content, "truncated": truncated, "size": size}

    if suffix == ".docx":
        return _docx_preview(target, size)

    mime, _ = mimetypes.guess_type(str(target))
    inline_ok = bool(mime) and (mime.startswith("image/") or mime == "application/pdf")
    if mime == "image/svg+xml":
        # SVG 可内嵌 <script>，浏览器同源渲染有 XSS 风险：只允许附件下载。
        inline_ok = False
    disposition = "inline" if inline_ok else "attachment"
    return FileResponse(
        path=str(target),
        media_type=mime or "application/octet-stream",
        headers={"Content-Disposition": f'{disposition}; filename="{target.name}"'},
    )


@router.post("/workspace/open")
async def open_in_system(request: Request, path: str = ""):
    """W4 用系统默认程序打开文件 / 用资源管理器定位目录。

    默认关闭：环境变量 ``RA_ALLOW_SHELL_OPEN`` 不为 ``"1"`` 时一律 403，
    防止 Web 模式下被滥用唤起系统 shell。子进程经 subprocess.Popen 启动，
    不等待其退出。
    """
    if os.getenv("RA_ALLOW_SHELL_OPEN", "").strip() != "1":
        raise HTTPException(status_code=403, detail="未启用：需设置环境变量 RA_ALLOW_SHELL_OPEN=1")
    root = _root_of(request)
    target = _fence(path, root)
    if not (target.is_file() or target.is_dir()):
        raise HTTPException(status_code=404, detail="目标不存在")

    system = platform.system()
    try:
        if system == "Windows":
            if target.is_dir():
                subprocess.Popen(["explorer", str(target)])
            else:
                os.startfile(str(target))
        elif system == "Darwin":
            subprocess.Popen(["open", str(target)])
        else:
            subprocess.Popen(["xdg-open", str(target)])
    except OSError as exc:
        raise HTTPException(status_code=500, detail=f"无法打开目标: {exc}") from exc
    return {"ok": True, "path": _rel_posix(target, root)}

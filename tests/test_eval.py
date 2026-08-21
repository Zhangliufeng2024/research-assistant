"""Tests for eval harness — metrics and thresholds (fully offline)."""

import json
from pathlib import Path

import pytest

yaml = pytest.importorskip("yaml")

from research_assistant.eval.metrics import (
    check_thresholds,
    compute_metrics,
)
from research_assistant.eval.runner import load_task


def _build_run_dir(tmp_path: Path) -> Path:
    """Synthesize a finished run directory with docx, gates report, run.json."""
    from docx import Document

    run_dir = tmp_path / "run"
    (run_dir / "final").mkdir(parents=True)
    (run_dir / "figures").mkdir()
    (run_dir / "drafts").mkdir()

    doc = Document()
    for s in ("Introduction", "Methods", "Results", "Discussion", "Conclusion"):
        doc.add_heading(s, level=1)
        doc.add_paragraph(" ".join(["word"] * 300))
    doc.save(str(run_dir / "final" / "manuscript.docx"))

    (run_dir / "figures" / "fig1.png").write_bytes(b"png")
    (run_dir / "figures" / "fig2.png").write_bytes(b"png")

    (run_dir / "gates_report.json").write_text(json.dumps({
        "passed": True,
        "revision_round": 0,
        "results": [
            {"name": "citations", "passed": True, "severity": "blocking",
             "failures": [],
             "details": {"total": 12, "verified": 11, "uncertain": 1,
                         "unverified": 0, "pass_rate": 1.0}},
            {"name": "document", "passed": True, "severity": "blocking",
             "failures": [], "details": {"word_count": 1503}},
        ],
    }), encoding="utf-8")

    (run_dir / "run.json").write_text(json.dumps({
        "status": "complete",
        "budget": {"cost_usd": 0.42, "total_tokens": 91_000, "turns": 37},
    }), encoding="utf-8")
    return run_dir


@pytest.fixture
def run_dir(tmp_path):
    return _build_run_dir(tmp_path)


class TestComputeMetrics:
    def test_all_metrics_present(self, run_dir):
        m = compute_metrics(run_dir)
        assert m["gates_passed"] is True
        assert m["has_final_docx"] is True
        assert m["figures_count"] == 2
        assert m["citations_count"] == 12
        assert m["citation_pass_rate"] == 1.0
        assert m["unverified"] == 0
        assert m["cost_usd"] == 0.42
        assert m["total_tokens"] == 91_000
        assert m["word_count"] > 1000

    def test_missing_run_dir_is_safe(self, tmp_path):
        m = compute_metrics(tmp_path / "nothing")
        assert m["gates_passed"] is None
        assert m["cost_usd"] is None


class TestThresholds:
    TASK = {
        "id": "t",
        "min_words": 900,
        "min_citations": 10,
        "citation_pass_rate": 0.9,
        "min_figures": 1,
        "budget_usd": 1.0,
    }

    def test_passing_task(self, run_dir):
        m = compute_metrics(run_dir)
        assert check_thresholds(m, self.TASK) == []

    def test_word_floor_violation(self, run_dir):
        m = compute_metrics(run_dir)
        failures = check_thresholds(m, {**self.TASK, "min_words": 99_999})
        assert any("word_count" in f for f in failures)

    def test_citation_floor_violation(self, run_dir):
        m = compute_metrics(run_dir)
        failures = check_thresholds(m, {**self.TASK, "min_citations": 50})
        assert any("citations_count" in f for f in failures)

    def test_budget_violation(self, run_dir):
        m = compute_metrics(run_dir)
        failures = check_thresholds(m, {**self.TASK, "budget_usd": 0.10})
        assert any("budget_usd" in f for f in failures)

    def test_missing_sections_surfaced_from_gate_report(self, run_dir):
        report = json.loads((run_dir / "gates_report.json").read_text(encoding="utf-8"))
        report["results"][1]["failures"] = ["missing sections: Discussion"]
        (run_dir / "gates_report.json").write_text(json.dumps(report))
        m = compute_metrics(run_dir)
        failures = check_thresholds(m, {**self.TASK,
                                        "expect_sections": ["Discussion"]})
        assert any("missing sections" in f for f in failures)


class TestLoadTask:
    def test_valid_task(self, tmp_path):
        p = tmp_path / "t.yaml"
        p.write_text(yaml.safe_dump({"id": "x", "query": "write", "min_words": 10}))
        task = load_task(p)
        assert task["id"] == "x"

    def test_missing_query_rejected(self, tmp_path):
        p = tmp_path / "t.yaml"
        p.write_text(yaml.safe_dump({"id": "x"}))
        with pytest.raises(ValueError, match="query"):
            load_task(p)

    def test_golden_tasks_in_repo_are_valid(self):
        tasks_dir = Path(__file__).parent.parent / "eval" / "golden_tasks"
        files = sorted(tasks_dir.glob("*.yaml"))
        assert len(files) >= 3
        for f in files:
            task = load_task(f)
            assert task.get("query", "").strip()
            assert task.get("budget_usd"), f"{f.name} should cap budget_usd"

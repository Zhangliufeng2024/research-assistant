"""Tests for research_assistant.docgen — PaperBuilder, BibTeX parsing, templates, cross-refs."""

from pathlib import Path

from docx import Document

from research_assistant.docgen import (
    CrossReferenceManager,
    DocumentTemplate,
    PaperBuilder,
    Reference,
    get_template,
    list_templates,
    parse_bibtex,
)


class TestReference:
    def test_format_apa_article(self):
        ref = Reference(
            key="smith2023",
            authors="Smith, J. and Doe, A.",
            title="A Great Paper",
            year=2023,
            journal="Nature",
            volume="42",
            pages="1-10",
            doi="10.1234/example",
        )
        text = ref.format_apa(1)
        assert text.startswith("[1]")
        assert "Smith, J. and Doe, A." in text
        assert '"A Great Paper."' in text
        assert "Nature," in text
        assert "vol. 42" in text
        assert "pp. 1-10" in text
        assert "2023." in text
        assert "DOI: 10.1234/example" in text

    def test_format_apa_inproceedings(self):
        ref = Reference(
            key="lee2022",
            authors="Lee, K.",
            title="Conference Paper",
            year=2022,
            booktitle="ICML 2022",
            pages="100-110",
        )
        text = ref.format_apa(3)
        assert "[3]" in text
        assert "in ICML 2022," in text

    def test_format_apa_minimal(self):
        ref = Reference(key="x", authors="A", title="T", year=2020)
        text = ref.format_apa(1)
        assert "[1]" in text
        assert "2020." in text


class TestPaperBuilder:
    def test_create_empty_doc(self, tmp_path):
        paper = PaperBuilder("Test Title")
        out = paper.save(str(tmp_path / "test.docx"))
        assert Path(out).exists()
        doc = Document(out)
        assert len(doc.paragraphs) > 0

    def test_title_and_authors(self, tmp_path):
        paper = PaperBuilder("My Paper", authors=["Alice", "Bob"])
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("My Paper" in t for t in texts)
        assert any("Alice, Bob" in t for t in texts)

    def test_abstract(self, tmp_path):
        paper = PaperBuilder("Title", abstract="This is the abstract text.")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Abstract" in t for t in texts)
        assert any("abstract text" in t for t in texts)

    def test_affiliation(self, tmp_path):
        paper = PaperBuilder("Title", authors=["A"], affiliation="MIT")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("MIT" in t for t in texts)

    def test_add_section(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_section("Introduction", "First paragraph.\n\nSecond paragraph.")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Introduction" in t for t in texts)
        assert any("First paragraph." in t for t in texts)
        assert any("Second paragraph." in t for t in texts)

    def test_add_section_levels(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_section("Level 1", "Content", level=1)
        paper.add_section("Level 2", "Content", level=2)
        paper.add_section("Level 3", "Content", level=3)
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 3

    def test_add_section_level_clamped(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_section("Too High", "Content", level=5)
        paper.add_section("Too Low", "Content", level=0)
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 2

    def test_add_figure_missing_image(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_figure(str(tmp_path / "nonexistent.png"), "A caption")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Missing figure" in t for t in texts)
        assert paper._figure_count == 0

    def test_add_figure_with_image(self, tmp_path):
        # Create a minimal valid PNG (1x1 pixel)
        import struct
        import zlib
        def _make_png():
            sig = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
            ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
            raw = b"\x00\x00\x00\x00"
            idat_data = zlib.compress(raw)
            idat_crc = zlib.crc32(b"IDAT" + idat_data) & 0xFFFFFFFF
            idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", idat_crc)
            iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
            iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
            return sig + ihdr + idat + iend

        img = tmp_path / "fig1.png"
        img.write_bytes(_make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img), "Test caption")
        assert paper._figure_count == 1
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Figure 1. Test caption" in t for t in texts)

    def test_figure_auto_numbering(self, tmp_path):
        import struct
        import zlib
        def _make_png():
            sig = b"\x89PNG\r\n\x1a\n"
            ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
            ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
            ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
            raw = b"\x00\x00\x00\x00"
            idat_data = zlib.compress(raw)
            idat_crc = zlib.crc32(b"IDAT" + idat_data) & 0xFFFFFFFF
            idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", idat_crc)
            iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
            iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
            return sig + ihdr + idat + iend

        img1 = tmp_path / "fig1.png"
        img2 = tmp_path / "fig2.png"
        img1.write_bytes(_make_png())
        img2.write_bytes(_make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img1), "First")
        paper.add_figure(str(img2), "Second")
        assert paper._figure_count == 2
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Figure 1." in t for t in texts)
        assert any("Figure 2." in t for t in texts)

    def test_add_table(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_table(["A", "B"], [["1", "2"], ["3", "4"]], caption="Results")
        assert paper._table_count == 1
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "A"
        texts = [p.text for p in doc.paragraphs]
        assert any("Table 1. Results" in t for t in texts)

    def test_table_auto_numbering(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_table(["X"], [["1"]], caption="First")
        paper.add_table(["Y"], [["2"]], caption="Second")
        assert paper._table_count == 2
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("Table 1." in t for t in texts)
        assert any("Table 2." in t for t in texts)

    def test_add_references(self, tmp_path):
        refs = [
            Reference(key="a", authors="Smith", title="Paper A", year=2020, journal="J1"),
            Reference(key="b", authors="Jones", title="Paper B", year=2021, journal="J2"),
        ]
        paper = PaperBuilder("Title")
        paper.add_references(refs)
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("References" in t for t in texts)
        assert any("[1]" in t for t in texts)
        assert any("[2]" in t for t in texts)

    def test_save_creates_parents(self, tmp_path):
        paper = PaperBuilder("Title")
        out = paper.save(str(tmp_path / "a" / "b" / "paper.docx"))
        assert Path(out).exists()


class TestParseBibtex:
    def test_nonexistent_file(self):
        refs = parse_bibtex("/nonexistent/path.bib")
        assert refs == []

    def test_empty_file(self, tmp_path):
        bib = tmp_path / "empty.bib"
        bib.write_text("")
        refs = parse_bibtex(str(bib))
        assert refs == []

    def test_single_article(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("""@article{smith2023,
  author = {Smith, John and Doe, Jane},
  title = {A Great Discovery},
  journal = {Nature},
  year = {2023},
  volume = {42},
  pages = {1-10},
  doi = {10.1234/example},
}
""")
        refs = parse_bibtex(str(bib))
        assert len(refs) == 1
        r = refs[0]
        assert r.key == "smith2023"
        assert r.authors == "Smith, John and Doe, Jane"
        assert r.title == "A Great Discovery"
        assert r.year == 2023
        assert r.journal == "Nature"
        assert r.volume == "42"
        assert r.pages == "1-10"
        assert r.doi == "10.1234/example"
        assert r.ref_type == "article"

    def test_multiple_entries(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("""@article{a,
  author = {A},
  title = {Title A},
  year = {2020},
}

@inproceedings{b,
  author = {B},
  title = {Title B},
  year = {2021},
  booktitle = {ICML},
}
""")
        refs = parse_bibtex(str(bib))
        assert len(refs) == 2
        assert refs[0].ref_type == "article"
        assert refs[1].ref_type == "inproceedings"
        assert refs[1].booktitle == "ICML"

    def test_invalid_year(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("""@article{x,
  author = {X},
  title = {Y},
  year = {unknown},
}
""")
        refs = parse_bibtex(str(bib))
        assert len(refs) == 1
        assert refs[0].year == 0

    def test_add_references_from_bibtex(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("""@article{test,
  author = {Test Author},
  title = {Test Title},
  journal = {Test Journal},
  year = {2024},
}
""")
        paper = PaperBuilder("Title")
        paper.add_references_from_bibtex(str(bib))
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        texts = [p.text for p in doc.paragraphs]
        assert any("References" in t for t in texts)
        assert any("[1]" in t and "Test Author" in t for t in texts)


# ---------------------------------------------------------------------------
# Template system
# ---------------------------------------------------------------------------

class TestTemplateRegistry:
    def test_list_templates_includes_defaults(self):
        names = list_templates()
        assert "default" in names
        assert "ieee_journal" in names
        assert "nature" in names
        assert "apa7" in names
        assert "asce" in names
        assert "elsevier" in names

    def test_get_template_known(self):
        t = get_template("ieee_journal")
        assert t.name == "ieee_journal"
        assert t.columns == 2

    def test_get_template_unknown_returns_default(self):
        t = get_template("nonexistent_venue")
        assert t.name == "default"

    def test_template_fields(self):
        t = get_template("nature")
        assert t.body_font == "Arial"
        assert t.heading_font == "Arial"
        assert t.columns == 1

    def test_apa7_double_spaced(self):
        t = get_template("apa7")
        assert t.line_spacing == 2.0


class TestPaperBuilderTemplate:
    def test_default_template(self, tmp_path):
        paper = PaperBuilder("Title")
        assert paper.template.name == "default"
        out = paper.save(str(tmp_path / "paper.docx"))
        assert Path(out).exists()

    def test_ieee_template_string(self, tmp_path):
        paper = PaperBuilder("IEEE Paper", template="ieee_journal")
        assert paper.template.name == "ieee_journal"
        paper.add_section("Introduction", "Content here.")
        out = paper.save(str(tmp_path / "ieee.docx"))
        assert Path(out).exists()

    def test_custom_template_object(self, tmp_path):
        custom = DocumentTemplate(
            name="custom",
            display_name="My Custom",
            body_font="Arial",
            body_font_size=14.0,
        )
        paper = PaperBuilder("Custom Paper", template=custom)
        assert paper.template.name == "custom"
        assert paper.template.body_font == "Arial"
        out = paper.save(str(tmp_path / "custom.docx"))
        assert Path(out).exists()

    def test_available_templates_static_method(self):
        templates = PaperBuilder.available_templates()
        assert isinstance(templates, list)
        assert "default" in templates

    def test_numbered_headings_ieee(self, tmp_path):
        paper = PaperBuilder("Title", template="ieee_journal")
        paper.add_section("Introduction", "Body.")
        paper.add_section("Methods", "Body.")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("1." in h for h in headings)
        assert any("2." in h for h in headings)

    def test_unnumbered_headings_nature(self, tmp_path):
        paper = PaperBuilder("Title", template="nature")
        paper.add_section("Introduction", "Body.")
        out = paper.save(str(tmp_path / "paper.docx"))
        doc = Document(out)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert any("Introduction" == h for h in headings)


# ---------------------------------------------------------------------------
# Cross-reference system
# ---------------------------------------------------------------------------

class TestCrossReferenceManager:
    def test_register_and_get(self):
        refs = CrossReferenceManager()
        refs.register("fig:overview", "Figure", 1)
        entry = refs.get("fig:overview")
        assert entry is not None
        assert entry.ref_type == "Figure"
        assert entry.number == 1

    def test_get_unknown_returns_none(self):
        refs = CrossReferenceManager()
        assert refs.get("nonexistent") is None

    def test_format_ref_known(self):
        refs = CrossReferenceManager()
        refs.register("tab:results", "Table", 3)
        assert refs.format_ref("tab:results") == "Table 3"

    def test_format_ref_unknown(self):
        refs = CrossReferenceManager()
        assert refs.format_ref("missing") == "[??missing]"

    def test_resolve_replaces_markers(self):
        refs = CrossReferenceManager()
        refs.register("fig:a", "Figure", 1)
        refs.register("tab:b", "Table", 2)
        text = "As shown in {ref:fig:a} and {ref:tab:b}, the results are clear."
        resolved = refs.resolve(text)
        assert resolved == "As shown in Figure 1 and Table 2, the results are clear."

    def test_resolve_unknown_label(self):
        refs = CrossReferenceManager()
        text = "See {ref:unknown}."
        resolved = refs.resolve(text)
        assert resolved == "See [??unknown]."

    def test_resolve_no_markers(self):
        refs = CrossReferenceManager()
        text = "Plain text without references."
        assert refs.resolve(text) == text


class TestPaperBuilderCrossRefs:
    def _make_png(self):
        import struct
        import zlib
        sig = b"\x89PNG\r\n\x1a\n"
        ihdr_data = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
        ihdr_crc = zlib.crc32(b"IHDR" + ihdr_data) & 0xFFFFFFFF
        ihdr = struct.pack(">I", 13) + b"IHDR" + ihdr_data + struct.pack(">I", ihdr_crc)
        raw = b"\x00\x00\x00\x00"
        idat_data = zlib.compress(raw)
        idat_crc = zlib.crc32(b"IDAT" + idat_data) & 0xFFFFFFFF
        idat = struct.pack(">I", len(idat_data)) + b"IDAT" + idat_data + struct.pack(">I", idat_crc)
        iend_crc = zlib.crc32(b"IEND") & 0xFFFFFFFF
        iend = struct.pack(">I", 0) + b"IEND" + struct.pack(">I", iend_crc)
        return sig + ihdr + idat + iend

    def test_figure_label_and_ref(self, tmp_path):
        img = tmp_path / "fig.png"
        img.write_bytes(self._make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img), "Overview", label="fig:overview")
        assert paper.ref("fig:overview") == "Figure 1"

    def test_table_label_and_ref(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_table(["A"], [["1"]], caption="Results", label="tab:results")
        assert paper.ref("tab:results") == "Table 1"

    def test_section_label_and_ref(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_section("Introduction", "Content", label="sec:intro")
        assert paper.ref("sec:intro") == "Section 1"

    def test_ref_unknown_label(self, tmp_path):
        paper = PaperBuilder("Title")
        assert paper.ref("nonexistent") == "[??nonexistent]"

    def test_cross_refs_resolved_in_saved_doc(self, tmp_path):
        img = tmp_path / "fig.png"
        img.write_bytes(self._make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img), "Architecture", label="fig:arch")
        paper.add_table(["X"], [["1"]], caption="Data", label="tab:data")
        paper.add_section("Discussion", "See {ref:fig:arch} and {ref:tab:data} for details.")
        out = paper.save(str(tmp_path / "paper.docx"))

        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Figure 1" in all_text
        assert "Table 1" in all_text
        assert "{ref:" not in all_text

    def test_cross_refs_in_tables(self, tmp_path):
        img = tmp_path / "fig.png"
        img.write_bytes(self._make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img), "Arch", label="fig:x")
        paper.add_table(["Col"], [["{ref:fig:x}"]], caption="Refs in table")
        out = paper.save(str(tmp_path / "paper.docx"))

        doc = Document(out)
        table_text = " ".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        assert "Figure 1" in table_text
        assert "{ref:" not in table_text

    def test_multiple_figures_numbering(self, tmp_path):
        img = tmp_path / "fig.png"
        img.write_bytes(self._make_png())

        paper = PaperBuilder("Title")
        paper.add_figure(str(img), "First", label="fig:a")
        paper.add_figure(str(img), "Second", label="fig:b")
        paper.add_section("Text", "Comparing {ref:fig:a} and {ref:fig:b}.")
        out = paper.save(str(tmp_path / "paper.docx"))

        doc = Document(out)
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "Figure 1" in all_text
        assert "Figure 2" in all_text

    def test_refs_property_access(self, tmp_path):
        paper = PaperBuilder("Title")
        paper.add_section("Intro", "Hi", label="sec:intro")
        entry = paper.refs.get("sec:intro")
        assert entry is not None
        assert entry.ref_type == "Section"


class TestReferenceFormatNamed:
    def test_format_numbered_is_format_apa(self):
        ref = Reference(key="x", authors="A", title="T", year=2020)
        assert ref.format_numbered(1) == ref.format_apa(1)

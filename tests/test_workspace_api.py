"""Tests for the workspace API (R5 计划 R1: W1-W4).

Covers: /api/workspace 名片、/api/workspace/tree 懒加载清单（忽略项过滤、
排序与条目形状、子目录下钻、safe_resolve 围栏穿越拒绝）、
/api/workspace/file 泛化预览（文本截断标志、无扩展名小文件、真实
python-docx 生成的 docx 抽取、PNG inline 头、未知二进制 attachment）、
/api/workspace/open 的 RA_ALLOW_SHELL_OPEN 开关。

App construction mirrors test_web_api.py: bare FastAPI + hand-wired
app.state（不跑真实 lifespan）；工作区根 = monkeypatch.chdir 后的 tmp_path。
"""

import base64
import os
import subprocess
import sys
from pathlib import Path

import pytest

fastapi = pytest.importorskip("fastapi")
from fastapi import FastAPI  # noqa: E402
from fastapi.testclient import TestClient  # noqa: E402

from research_assistant.web.workspace import (  # noqa: E402
    DOCX_PARAGRAPH_LIMIT,
    TEXT_PREVIEW_LIMIT,
)
from research_assistant.web.workspace import (  # noqa: E402
    router as workspace_router,
)

#: 标准 1x1 PNG（base64），用于验证图片 inline 预览分支。
PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJ"
    "AAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)


def _client(tmp_path: Path, monkeypatch: pytest.MonkeyPatch) -> TestClient:
    """裸 app + 手工接线 state；工作区根即 tmp_path。"""
    monkeypatch.chdir(tmp_path)
    app = FastAPI()
    app.include_router(workspace_router, prefix="/api")
    app.state.output_folder = tmp_path / "writing_outputs"
    return TestClient(app)


def _patch_launchers(monkeypatch: pytest.MonkeyPatch) -> list[list[str]]:
    """把 explorer/open/xdg-open/startfile 全部替换为记录器，返回调用列表。"""
    calls: list[list[str]] = []

    def fake_popen(argv, *args, **kwargs):
        calls.append([str(a) for a in argv])

    monkeypatch.setattr(subprocess, "Popen", fake_popen)
    # os.startfile 仅 Windows 存在；raising=False 允许在其它平台打桩。
    monkeypatch.setattr(
        os, "startfile", lambda p: calls.append(["startfile", str(p)]),
        raising=False,
    )
    return calls


# ---------------------------------------------------------------------------
# W1: GET /api/workspace
# ---------------------------------------------------------------------------

class TestWorkspaceCard:
    def test_card_fields(self, tmp_path, monkeypatch):
        (tmp_path / ".git").mkdir()
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace")
        assert resp.status_code == 200
        card = resp.json()
        assert Path(card["root"]).resolve() == tmp_path.resolve()
        assert card["name"] == tmp_path.name
        assert card["output_folder"] == "writing_outputs"
        assert card["has_git"] is True

    def test_card_without_git(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        card = client.get("/api/workspace").json()
        assert card["has_git"] is False

    def test_output_folder_outside_root_falls_back_to_absolute(self, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        app = FastAPI()
        app.include_router(workspace_router, prefix="/api")
        external = tmp_path.parent / "somewhere_else"
        app.state.output_folder = external
        client = TestClient(app)
        card = client.get("/api/workspace").json()
        assert Path(card["output_folder"]) == external

    def test_card_tolerates_missing_state(self, tmp_path, monkeypatch):
        monkeypatch.chdir(tmp_path)
        app = FastAPI()
        app.include_router(workspace_router, prefix="/api")
        client = TestClient(app)
        card = client.get("/api/workspace").json()
        assert card["output_folder"] is None


# ---------------------------------------------------------------------------
# W2: GET /api/workspace/tree
# ---------------------------------------------------------------------------

class TestTree:
    def test_ignore_list_and_hidden_filtering(self, tmp_path, monkeypatch):
        for name in (".git", "__pycache__", ".ra", "node_modules"):
            (tmp_path / name).mkdir()
        (tmp_path / ".hidden").write_text("x", encoding="utf-8")
        (tmp_path / "src").mkdir()
        (tmp_path / "README.md").write_text("# hi", encoding="utf-8")

        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/tree").json()

        names = [item["name"] for item in data["items"]]
        assert names == ["src", "README.md"]  # 内部目录与隐藏项全部被滤掉

    def test_shape_and_sorting_dirs_first(self, tmp_path, monkeypatch):
        (tmp_path / "b_dir").mkdir()
        (tmp_path / "a_dir").mkdir()
        (tmp_path / "z.txt").write_text("z", encoding="utf-8")
        (tmp_path / "A.txt").write_text("A", encoding="utf-8")

        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/tree").json()

        assert data["path"] == ""
        assert [i["name"] for i in data["items"]] == ["a_dir", "b_dir", "A.txt", "z.txt"]
        by_name = {i["name"]: i for i in data["items"]}
        for item in data["items"]:
            assert set(item) == {"name", "path", "type", "size", "mtime"}
            assert "\\" not in item["path"]  # 相对 root 的 POSIX 路径
            assert isinstance(item["mtime"], float)
        assert by_name["b_dir"]["type"] == "dir"
        assert by_name["b_dir"]["path"] == "b_dir"
        assert by_name["b_dir"]["size"] is None
        assert by_name["A.txt"]["type"] == "file"
        assert by_name["A.txt"]["size"] == 1
        assert isinstance(by_name["A.txt"]["mtime"], float)

    def test_subdirectory_lazy_load(self, tmp_path, monkeypatch):
        sub = tmp_path / "src" / "deep"
        sub.mkdir(parents=True)
        (sub / "util.py").write_text("print(1)", encoding="utf-8")

        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/tree", params={"path": "src/deep"}).json()

        assert data["path"] == "src/deep"
        assert [(i["name"], i["type"]) for i in data["items"]] == [("util.py", "file")]

    def test_missing_dir_404(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        assert client.get("/api/workspace/tree", params={"path": "nope"}).status_code == 404

    def test_file_target_404(self, tmp_path, monkeypatch):
        (tmp_path / "f.txt").write_text("x", encoding="utf-8")
        client = _client(tmp_path, monkeypatch)
        assert client.get("/api/workspace/tree", params={"path": "f.txt"}).status_code == 404


class TestFence:
    def test_tree_dotdot_traversal_403(self, tmp_path, monkeypatch):
        outside = tmp_path.parent
        (outside / "secret.txt").write_text("top secret", encoding="utf-8")
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/tree", params={"path": "../../"})
        assert resp.status_code == 403
        assert "secret" not in resp.text

    def test_tree_url_encoded_traversal_403(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/tree?path=%2E%2E%2F%2E%2E")
        assert resp.status_code == 403

    def test_tree_absolute_path_escape_403(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        absolute = str(tmp_path.parent / "elsewhere")
        resp = client.get("/api/workspace/tree", params={"path": absolute})
        assert resp.status_code == 403

    def test_file_traversal_403(self, tmp_path, monkeypatch):
        (tmp_path.parent / "outside.bin").write_bytes(b"\x00\x01")
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "../outside.bin"})
        assert resp.status_code == 403


# ---------------------------------------------------------------------------
# W3: GET /api/workspace/file
# ---------------------------------------------------------------------------

class TestFilePreviewText:
    def test_small_text(self, tmp_path, monkeypatch):
        # write_bytes 避免 Windows 文本模式把 \n 翻译成 \r\n
        (tmp_path / "notes.md").write_bytes("# 标题\n正文".encode())
        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/file", params={"path": "notes.md"}).json()
        assert data == {
            "kind": "text",
            "content": "# 标题\n正文",
            "truncated": False,
            "size": len("# 标题\n正文".encode()),
        }

    def test_large_text_truncated_flag(self, tmp_path, monkeypatch):
        payload = b"A" * (TEXT_PREVIEW_LIMIT + 10)
        (tmp_path / "big.log").write_bytes(payload)
        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/file", params={"path": "big.log"}).json()
        assert data["kind"] == "text"
        assert data["truncated"] is True
        assert data["size"] == TEXT_PREVIEW_LIMIT + 10
        assert len(data["content"]) == TEXT_PREVIEW_LIMIT  # 只回头部，不整读大文件

    def test_extensionless_small_file_is_text(self, tmp_path, monkeypatch):
        (tmp_path / "Makefile").write_text("all:\n\techo hi\n", encoding="utf-8")
        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/file", params={"path": "Makefile"}).json()
        assert data["kind"] == "text"
        assert data["truncated"] is False
        assert "echo hi" in data["content"]

    def test_invalid_utf8_replaced(self, tmp_path, monkeypatch):
        (tmp_path / "mixed.txt").write_bytes(b"ok\xff\xfe")
        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/file", params={"path": "mixed.txt"}).json()
        assert data["kind"] == "text"
        assert data["content"].startswith("ok")


class TestFilePreviewDocx:
    def test_docx_paragraph_and_table_extraction(self, tmp_path, monkeypatch):
        from docx import Document  # python-docx 为项目硬依赖，直接使用

        document = Document()
        for i in range(DOCX_PARAGRAPH_LIMIT + 50):  # 超出抽取上限，验证截取
            document.add_paragraph(f"P{i}")
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "CELL_A"
        table.cell(1, 1).text = "CELL_B"
        target = tmp_path / "report.docx"
        document.save(str(target))

        client = _client(tmp_path, monkeypatch)
        data = client.get("/api/workspace/file", params={"path": "report.docx"}).json()

        assert data["kind"] == "text"
        assert data["truncated"] is False
        assert data["size"] == target.stat().st_size  # size 报源文件大小
        assert "P0" in data["content"]
        assert f"P{DOCX_PARAGRAPH_LIMIT - 1}" in data["content"]
        assert f"P{DOCX_PARAGRAPH_LIMIT}" not in data["content"]  # 第 200 段之后不抽
        assert "CELL_A" in data["content"]
        assert "CELL_B" in data["content"]

    def test_corrupt_docx_falls_back_to_attachment(self, tmp_path, monkeypatch):
        (tmp_path / "broken.docx").write_bytes(b"not a zip")
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "broken.docx"})
        assert resp.status_code == 200
        assert resp.headers["content-disposition"].startswith("attachment")


class TestFilePreviewBinary:
    def test_png_served_inline(self, tmp_path, monkeypatch):
        (tmp_path / "figure.png").write_bytes(PNG_1PX)
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "figure.png"})
        assert resp.status_code == 200
        assert resp.headers["content-type"] == "image/png"
        assert resp.headers["content-disposition"].startswith("inline")

    def test_unknown_binary_served_as_attachment(self, tmp_path, monkeypatch):
        (tmp_path / "model.pkl").write_bytes(b"\x93PK")
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "model.pkl"})
        assert resp.status_code == 200
        assert resp.headers["content-disposition"].startswith("attachment")
        assert resp.headers["content-type"] == "application/octet-stream"

    def test_zip_gets_attachment_disposition(self, tmp_path, monkeypatch):
        (tmp_path / "bundle.zip").write_bytes(b"PK\x03\x04")
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "bundle.zip"})
        assert resp.headers["content-disposition"].startswith("attachment")

    def test_missing_file_404(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "ghost.txt"})
        assert resp.status_code == 404

    def test_directory_target_404(self, tmp_path, monkeypatch):
        (tmp_path / "adir").mkdir()
        client = _client(tmp_path, monkeypatch)
        resp = client.get("/api/workspace/file", params={"path": "adir"})
        assert resp.status_code == 404


# ---------------------------------------------------------------------------
# W4: POST /api/workspace/open
# ---------------------------------------------------------------------------

class TestOpen:
    def test_disabled_by_default_403(self, tmp_path, monkeypatch):
        monkeypatch.delenv("RA_ALLOW_SHELL_OPEN", raising=False)
        client = _client(tmp_path, monkeypatch)
        resp = client.post("/api/workspace/open", params={"path": ""})
        assert resp.status_code == 403

    def test_other_value_still_403(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "yes")
        client = _client(tmp_path, monkeypatch)
        assert client.post("/api/workspace/open", params={"path": ""}).status_code == 403

    @pytest.mark.skipif(sys.platform != "win32", reason="explorer/startfile 分支")
    def test_windows_dir_locates_via_explorer(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "1")
        calls = _patch_launchers(monkeypatch)
        (tmp_path / "paper").mkdir()
        client = _client(tmp_path, monkeypatch)

        resp = client.post("/api/workspace/open", params={"path": "paper"})

        assert resp.status_code == 200
        assert resp.json() == {"ok": True, "path": "paper"}
        assert calls[0][0] == "explorer"
        assert Path(calls[0][1]).resolve() == (tmp_path / "paper").resolve()

    @pytest.mark.skipif(sys.platform != "win32", reason="os.startfile 分支")
    def test_windows_file_opens_via_startfile(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "1")
        calls = _patch_launchers(monkeypatch)
        (tmp_path / "doc.txt").write_text("hi", encoding="utf-8")
        client = _client(tmp_path, monkeypatch)

        resp = client.post("/api/workspace/open", params={"path": "doc.txt"})

        assert resp.status_code == 200
        assert calls[0][0] == "startfile"
        assert Path(calls[0][1]).resolve() == (tmp_path / "doc.txt").resolve()

    @pytest.mark.skipif(sys.platform == "win32", reason="非 Windows 走 Popen 分支")
    def test_posix_open_command(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "1")
        calls = _patch_launchers(monkeypatch)
        (tmp_path / "paper").mkdir()
        client = _client(tmp_path, monkeypatch)

        resp = client.post("/api/workspace/open", params={"path": "paper"})

        assert resp.status_code == 200
        expected = "open" if sys.platform == "darwin" else "xdg-open"
        assert calls[0][0] == expected
        assert Path(calls[0][-1]).resolve() == (tmp_path / "paper").resolve()

    def test_enabled_but_missing_target_404(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "1")
        _patch_launchers(monkeypatch)
        client = _client(tmp_path, monkeypatch)
        resp = client.post("/api/workspace/open", params={"path": "ghost"})
        assert resp.status_code == 404

    def test_enabled_traversal_still_fenced(self, tmp_path, monkeypatch):
        monkeypatch.setenv("RA_ALLOW_SHELL_OPEN", "1")
        calls = _patch_launchers(monkeypatch)
        client = _client(tmp_path, monkeypatch)
        resp = client.post("/api/workspace/open", params={"path": "../../"})
        assert resp.status_code == 403
        assert calls == []  # 围栏拦截在前，不会触发任何系统调用


# ---------------------------------------------------------------------------
# A+ 阶段 1 / F-1：快照缺失时恢复必须报 409，而不是删掉用户文件
#
# 旧实现：restore() 读不到快照就走「删除目标文件」分支，于是 Janitor 淘汰过
# bin 的老工作区里，用户在变更页点「恢复」= 销毁当前文件。
# ---------------------------------------------------------------------------

class TestRestoreMissingSnapshotIsRefused:
    def _make_change(self, tmp_path) -> tuple[str, Path]:
        from research_assistant.artifacts import ArtifactVersionStore

        target = tmp_path / "report.md"
        target.write_text("原始", encoding="utf-8")
        store = ArtifactVersionStore(tmp_path)
        rec = store.record(target, "原始".encode(), "改后".encode(), tool="write_file")
        assert rec is not None
        return rec["id"], target

    def test_restore_with_evicted_snapshot_returns_409_and_keeps_file(
        self, tmp_path, monkeypatch,
    ):
        from research_assistant.artifacts import ArtifactVersionStore

        change_id, target = self._make_change(tmp_path)
        ArtifactVersionStore(tmp_path).discard_snapshot(change_id, "before")

        client = _client(tmp_path, monkeypatch)
        resp = client.post(
            f"/api/workspace/changes/{change_id}/restore", json={"side": "before"},
        )

        assert resp.status_code == 409, resp.text
        assert "快照" in resp.json()["detail"]
        # 最关键：文件必须原封不动
        assert target.read_text(encoding="utf-8") == "原始"

    def test_normal_restore_still_works(self, tmp_path, monkeypatch):
        """反向断言：快照健在时恢复照常成功，不能被修 P0 时一并堵死。"""
        change_id, target = self._make_change(tmp_path)
        target.write_text("改后", encoding="utf-8")

        client = _client(tmp_path, monkeypatch)
        resp = client.post(
            f"/api/workspace/changes/{change_id}/restore", json={"side": "before"},
        )

        assert resp.status_code == 200, resp.text
        assert target.read_text(encoding="utf-8") == "原始"

    def test_unknown_change_returns_404(self, tmp_path, monkeypatch):
        client = _client(tmp_path, monkeypatch)
        resp = client.post(
            "/api/workspace/changes/does-not-exist/restore", json={"side": "before"},
        )
        assert resp.status_code == 404

    def test_bad_side_returns_422(self, tmp_path, monkeypatch):
        change_id, _target = self._make_change(tmp_path)
        client = _client(tmp_path, monkeypatch)
        resp = client.post(
            f"/api/workspace/changes/{change_id}/restore", json={"side": "middle"},
        )
        assert resp.status_code == 422

    def test_diff_reports_snapshot_availability(self, tmp_path, monkeypatch):
        """diff 结果带上可用性标记，供前端禁用已失效的恢复入口。"""
        from research_assistant.artifacts import ArtifactVersionStore

        change_id, _target = self._make_change(tmp_path)
        ArtifactVersionStore(tmp_path).discard_snapshot(change_id, "after")

        client = _client(tmp_path, monkeypatch)
        body = client.get(f"/api/workspace/changes/{change_id}").json()

        assert body["after_available"] is False
        assert body["before_available"] is True

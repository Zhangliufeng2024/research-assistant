"""Tests for research_assistant.gates — citation and document gates."""

import pytest

from research_assistant.gates import (
    CitationGate,
    DocGate,
    GateReport,
    GateResult,
    render_gate_report_md,
)
from research_assistant.tools.citation_verify import CitationResult, VerificationReport


def _report(*statuses: str) -> VerificationReport:
    rep = VerificationReport(total=len(statuses))
    for i, s in enumerate(statuses):
        rep.results.append(CitationResult(
            key=f"key{i}", title=f"Paper {i}", status=s,
            confidence=0.95 if s == "verified" else 0.3,
            doi="", year=2024, authors="Author",
        ))
        if s == "verified":
            rep.verified += 1
        elif s == "uncertain":
            rep.uncertain += 1
        else:
            rep.unverified += 1
    return rep


class TestCitationGate:
    @pytest.mark.asyncio
    async def test_all_verified_passes(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("@article{a, title={t}}")
        gate = CitationGate(bib, verify_fn=lambda p: _report("verified", "verified"))
        result = await gate.run({})
        assert result.passed
        assert result.details["unverified"] == 0

    @pytest.mark.asyncio
    async def test_unverified_blocks(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("@article{a, title={t}}")
        gate = CitationGate(bib, verify_fn=lambda p: _report("verified", "unverified"))
        result = await gate.run({})
        assert not result.passed
        assert any("UNVERIFIED [key1]" in f for f in result.failures)

    @pytest.mark.asyncio
    async def test_low_pass_rate_blocks(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("@article{a, title={t}}")
        gate = CitationGate(
            bib, min_pass_rate=0.95,
            verify_fn=lambda p: _report("uncertain", "verified", "verified", "verified"),
        )
        result = await gate.run({})
        # no unverified, but pass rate 100%? uncertain counts toward pass_rate -> passes
        # make it stricter: 3/4 = 75% < 95%
        gate2 = CitationGate(
            bib, min_pass_rate=0.95,
            verify_fn=lambda p: _report("uncertain", "uncertain", "uncertain", "verified"),
        )
        result2 = await gate2.run({})
        assert not result2.passed

    @pytest.mark.asyncio
    async def test_missing_bib_fails(self, tmp_path):
        gate = CitationGate(tmp_path / "nope.bib")
        result = await gate.run({})
        assert not result.passed
        assert "not found" in result.failures[0]

    @pytest.mark.asyncio
    async def test_report_saved(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("@article{a, title={t}}")
        out = tmp_path / "sources" / "CITATION_VERIFICATION.md"
        gate = CitationGate(
            bib, report_output=out,
            verify_fn=lambda p: _report("verified"),
        )
        await gate.run({})
        assert out.exists()


class TestGateReport:
    def test_passed_ignores_warn_severity(self):
        report = GateReport(results=[
            GateResult(name="a", passed=True),
            GateResult(name="b", passed=False, severity="warn", failures=["minor"]),
        ])
        assert report.passed
        assert report.blocking_failures == []

    def test_blocking_failure_detected(self):
        report = GateReport(results=[GateResult(name="a", passed=False, failures=["x"])])
        assert not report.passed
        assert "[a] x" in report.blocking_failures

    def test_save_load_roundtrip(self, tmp_path):
        report = GateReport(results=[
            GateResult(name="citations", passed=False, failures=["bad key"],
                       details={"total": 3}),
        ], revision_round=2)
        path = tmp_path / "gates_report.json"
        report.save(path)
        loaded = GateReport.load(path)
        assert loaded.revision_round == 2
        assert not loaded.passed
        assert loaded.results[0].details["total"] == 3

    def test_render_md_lists_failures(self):
        report = GateReport(results=[
            GateResult(name="doc", passed=False, failures=["missing sections: Methods"]),
        ])
        md = render_gate_report_md(report)
        assert "FAIL" in md
        assert "missing sections: Methods" in md


def _make_docx(path, sections, body_words=50, figure_refs=(), placeholder=None):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Title of the Paper")
    for s in sections:
        doc.add_heading(s, level=1)
        doc.add_paragraph(" ".join(["word"] * body_words))
    for ref in figure_refs:
        doc.add_paragraph(f"See figures/{ref} for details.")
    if placeholder:
        doc.add_paragraph(placeholder)
    doc.save(str(path))


class TestDocGate:
    @pytest.mark.asyncio
    async def test_complete_document_passes(self, tmp_path):
        docx = tmp_path / "manuscript.docx"
        _make_docx(docx, ["Introduction", "Methods", "Results"],
                   figure_refs=["fig1.png"])
        figs = tmp_path / "figures"
        figs.mkdir()
        (figs / "fig1.png").write_bytes(b"png")

        gate = DocGate(docx, expected_sections=["Introduction", "Methods", "Results"],
                       figures_dir=figs, target_words=150)
        result = await gate.run({})
        assert result.passed, result.failures

    @pytest.mark.asyncio
    async def test_missing_section_fails(self, tmp_path):
        docx = tmp_path / "m.docx"
        _make_docx(docx, ["Introduction"])
        gate = DocGate(docx, expected_sections=["Introduction", "Discussion"])
        result = await gate.run({})
        assert not result.passed
        assert any("Discussion" in f for f in result.failures)

    @pytest.mark.asyncio
    async def test_placeholder_fails(self, tmp_path):
        docx = tmp_path / "m.docx"
        _make_docx(docx, ["Introduction"], placeholder="TODO: add more citations")
        gate = DocGate(docx, expected_sections=["Introduction"])
        result = await gate.run({})
        assert not result.passed
        assert any("placeholder" in f.lower() for f in result.failures)

    @pytest.mark.asyncio
    async def test_missing_figure_file_fails(self, tmp_path):
        docx = tmp_path / "m.docx"
        _make_docx(docx, ["Introduction"], figure_refs=["ghost.png"])
        figs = tmp_path / "figures"
        figs.mkdir()
        gate = DocGate(docx, expected_sections=["Introduction"], figures_dir=figs)
        result = await gate.run({})
        assert not result.passed
        assert any("ghost.png" in f for f in result.failures)

    @pytest.mark.asyncio
    async def test_word_floor_fails(self, tmp_path):
        docx = tmp_path / "m.docx"
        _make_docx(docx, ["Introduction"], body_words=20)
        gate = DocGate(docx, expected_sections=["Introduction"], target_words=5_000)
        result = await gate.run({})
        assert not result.passed
        assert any("below floor" in f for f in result.failures)

    @pytest.mark.asyncio
    async def test_missing_file_fails(self, tmp_path):
        gate = DocGate(tmp_path / "ghost.docx")
        result = await gate.run({})
        assert not result.passed

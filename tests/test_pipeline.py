"""Tests for pipeline: artifacts, session store, and the state-machine runner.

The runner tests monkeypatch the sub-agent entry point so no LLM or network
is involved; the citation gate is stubbed to an all-verified report.
"""

import json
import re
from pathlib import Path

import pytest

from research_assistant.pipeline.artifacts import ArtifactStore
from research_assistant.session.store import SessionStore

# ---------------------------------------------------------------------------
# ArtifactStore
# ---------------------------------------------------------------------------

class TestArtifactStore:
    def test_register_and_valid(self, tmp_path):
        store = ArtifactStore(tmp_path)
        f = tmp_path / "plan.json"
        f.write_text('{"a": 1}')
        store.register("plan.json", f, stage="plan")
        assert store.is_valid("plan.json")

    def test_content_change_invalidates(self, tmp_path):
        store = ArtifactStore(tmp_path)
        f = tmp_path / "plan.json"
        f.write_text("v1")
        store.register("k", f, stage="plan")
        f.write_text("v2")
        assert not store.is_valid("k")

    def test_missing_file_invalid(self, tmp_path):
        store = ArtifactStore(tmp_path)
        f = tmp_path / "gone.txt"
        f.write_text("x")
        store.register("k", f, stage="s")
        f.unlink()
        assert not store.is_valid("k")

    def test_manifest_persists_across_instances(self, tmp_path):
        f = tmp_path / "a.txt"
        f.write_text("data")
        ArtifactStore(tmp_path).register("a", f, stage="s")
        again = ArtifactStore(tmp_path)
        assert again.is_valid("a")
        assert again.get("a").stage == "s"

    def test_all_valid(self, tmp_path):
        store = ArtifactStore(tmp_path)
        a, b = tmp_path / "a", tmp_path / "b"
        a.write_text("1")
        b.write_text("2")
        store.register("a", a, stage="s")
        store.register("b", b, stage="s")
        assert store.all_valid("a", "b")
        b.write_text("changed")
        assert not store.all_valid("a", "b")

    def test_manifest_replace_failure_preserves_previous_json(
            self, tmp_path, monkeypatch):
        first = tmp_path / "first.txt"
        first.write_text("one")
        store = ArtifactStore(tmp_path)
        store.register("first", first, stage="s")
        manifest = tmp_path / ".ra" / "artifacts" / "manifest.json"
        previous = manifest.read_text(encoding="utf-8")

        second = tmp_path / "second.txt"
        second.write_text("two")
        monkeypatch.setattr(
            "research_assistant.core.os.replace",
            lambda *args: (_ for _ in ()).throw(OSError("replace failed")),
        )
        with pytest.raises(OSError, match="replace failed"):
            store.register("second", second, stage="s")

        assert manifest.read_text(encoding="utf-8") == previous
        assert not list(manifest.parent.glob(".manifest.json.*.tmp"))


# ---------------------------------------------------------------------------
# SessionStore
# ---------------------------------------------------------------------------

class TestSessionStore:
    def test_create_and_reload(self, tmp_path):
        SessionStore.create(tmp_path, query="q", model="m")
        reloaded = SessionStore(tmp_path)
        assert reloaded.state.query == "q"
        assert reloaded.state.model == "m"
        assert reloaded.state.status == "running"

    def test_mark_stage_lifecycle(self, tmp_path):
        store = SessionStore.create(tmp_path, query="q", model="m")
        store.mark_stage("plan", "running")
        store.mark_stage("plan", "done", artifacts=["plan.json"])
        assert store.stage_status("plan") == "done"
        assert store.done_stages() == ["plan"]
        rec = store.state.stages["plan"]
        assert rec["finished_at"] > 0
        assert rec["artifacts"] == ["plan.json"]

    def test_events_append_and_read(self, tmp_path):
        store = SessionStore.create(tmp_path, query="q", model="m")
        store.log_event("stage_plan", {"n": 1})
        store.log_event("gates", {"passed": True})
        events = SessionStore(tmp_path).read_events()
        assert [e["kind"] for e in events] == ["stage_plan", "gates"]
        assert events[0]["data"]["n"] == 1

    def test_finish(self, tmp_path):
        store = SessionStore.create(tmp_path, query="q", model="m")
        store.finish("complete", budget.snapshot() if (budget := None) else {"cost_usd": 0.5})
        reloaded = SessionStore(tmp_path)
        assert reloaded.state.status == "complete"
        assert reloaded.state.budget["cost_usd"] == 0.5

    def test_corrupt_run_json_tolerated(self, tmp_path):
        tmp_path.mkdir(exist_ok=True)
        (tmp_path / "run.json").write_text("{not json", encoding="utf-8")
        store = SessionStore(tmp_path)
        assert store.state.status == "running"  # fresh state, no crash

    def test_save_replace_failure_preserves_previous_json(
            self, tmp_path, monkeypatch):
        store = SessionStore.create(tmp_path, query="original", model="m")
        run_file = tmp_path / "run.json"
        previous = run_file.read_text(encoding="utf-8")
        store.state.query = "changed"
        monkeypatch.setattr(
            "research_assistant.core.os.replace",
            lambda *args: (_ for _ in ()).throw(OSError("replace failed")),
        )
        with pytest.raises(OSError, match="replace failed"):
            store.save()

        assert run_file.read_text(encoding="utf-8") == previous
        assert not list(tmp_path.glob(".run.json.*.tmp"))


# ---------------------------------------------------------------------------
# Runner (sub-agents faked)
# ---------------------------------------------------------------------------

PLAN_JSON = {
    "paper_title": "Test Paper on FEM",
    "paper_type": "research_paper",
    "venue": "ASCE",
    "sections": [
        {"name": "introduction", "title": "Introduction",
         "description": "intro", "estimated_words": 100,
         "search_queries": ["fem"]},
        {"name": "methods", "title": "Methods",
         "description": "methods", "estimated_words": 100,
         "search_queries": ["fem"]},
    ],
    "figures": [],
    "total_estimated_words": 200,
    "citation_target": 4,
}


def _extract(prompt: str, pattern: str) -> str:
    m = re.search(pattern, prompt)
    return m.group(1).strip() if m else ""


def _make_fake_agent(calls: list, workdir_holder: dict):
    """Build a fake _run_stage_agent that creates the files each stage promises."""

    async def fake_agent(stage_name, prompt, system_prompt, *, model, work_dir,
                         api_key, base_url, provider, budget, hooks, cancel_event,
                         auto_continue=True, max_continuations=10,
                         approver=None, session_log=None, steer_queue=None,
                         write_anchor=None):  # R12 B5：任务侧写入归巢（fake 忽略）
        calls.append(stage_name)
        from research_assistant.agent import AgentResult

        if stage_name == "plan":
            return AgentResult(text_output=json.dumps(PLAN_JSON), stop_reason="completed")

        if stage_name.startswith("research_"):
            bib = Path(_extract(prompt, r"Write BibTeX entries to: (.+)"))
            summary = Path(_extract(prompt, r"Write a summary to: (.+)"))
            bib.parent.mkdir(parents=True, exist_ok=True)
            bib.write_text(
                "@article{real2024, author={A}, title={Real Paper},\n"
                " journal={J}, year={2024}}\n"
            )
            summary.write_text(f"Summary for {stage_name}\nkey findings here")
            return AgentResult(stop_reason="completed")

        if stage_name.startswith("figure_"):
            out = Path(_extract(prompt, r"Output: (.+)"))
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_bytes(b"png")
            return AgentResult(stop_reason="completed")

        if stage_name == "assemble":
            docx = Path(_extract(prompt, r"Then create the \.docx at: (.+)"))
            docx.parent.mkdir(parents=True, exist_ok=True)
            from docx import Document
            doc = Document()
            for s in PLAN_JSON["sections"]:
                doc.add_heading(s["title"], level=1)
                doc.add_paragraph(" ".join(["data"] * 150))
            doc.save(str(docx))
            # merge bib as the real agent would
            bib_file = Path(_extract(prompt, r"ONE file: (.+)"))
            bib_file.parent.mkdir(parents=True, exist_ok=True)
            if not bib_file.exists():
                srcs = sorted(bib_file.parent.parent.glob("sources/bib_*.bib"))
                bib_file.write_text("\n".join(
                    p.read_text(encoding="utf-8") for p in srcs))
            return AgentResult(stop_reason="completed")

        if stage_name.startswith("revision_"):
            raise AssertionError("gates should pass; revision must not run")

        if stage_name == "review":
            out_dir = Path(_extract(prompt, r"Write (.+)/PEER_REVIEW\.md"))
            (out_dir / "PEER_REVIEW.md").write_text("# review")
            (out_dir / "SUMMARY.md").write_text("# summary")
            return AgentResult(stop_reason="completed")

        raise AssertionError(f"unexpected stage {stage_name}")

    return fake_agent


@pytest.fixture
def gate_stub(monkeypatch):
    """CitationGate that always passes without network."""
    import research_assistant.gates as gates_pkg
    from research_assistant.gates import GateResult

    class StubCitationGate(gates_pkg.CitationGate):
        async def run(self, context):
            return GateResult(name="citations", passed=True,
                              details={"stubbed": True})

    monkeypatch.setattr(gates_pkg, "CitationGate", StubCitationGate)


async def _collect(agen):
    outs = []
    async for u in agen:
        outs.append(u)
    return outs


@pytest.mark.asyncio
async def test_full_pipeline_happy_path(tmp_path, monkeypatch, gate_stub):
    calls: list = []
    monkeypatch.setattr(
        "research_assistant.pipeline.runner._run_stage_agent",
        _make_fake_agent(calls, {}),
    )
    from research_assistant.pipeline.runner import run_pipeline

    out_dir = tmp_path / "writing_outputs" / "run1"
    updates = await _collect(run_pipeline(
        query="write a fem paper", model="test-model",
        work_dir=tmp_path, output_dir=out_dir,
    ))

    types = [u.get("type") for u in updates]
    assert "result" in types
    final = next(u for u in updates if u["type"] == "result")
    assert final["status"] == "success"
    assert (out_dir / "final" / "manuscript.docx").exists()
    assert (out_dir / "gates_report.json").exists()

    session = SessionStore(out_dir)
    assert session.state.status == "complete"
    assert session.stage_status("assemble") == "done"
    kinds = [e["kind"] for e in session.read_events()]
    assert "run_end" in kinds

    report = json.loads((out_dir / "gates_report.json").read_text(encoding="utf-8"))
    assert report["passed"] is True
    assert "revision_1" not in calls


@pytest.mark.asyncio
async def test_blocking_gate_failure_does_not_publish_final(
        tmp_path, monkeypatch, gate_stub):
    """Exhausted blocking gates must leave only drafts, never a final file."""
    import research_assistant.gates as gates_pkg
    from research_assistant.agent import AgentResult
    from research_assistant.gates import GateResult

    calls: list = []
    base_agent = _make_fake_agent(calls, {})

    async def fake_agent(stage_name, prompt, system_prompt, **kwargs):
        if stage_name.startswith("revision_"):
            calls.append(stage_name)
            return AgentResult(stop_reason="completed")
        return await base_agent(stage_name, prompt, system_prompt, **kwargs)

    class FailingDocGate(gates_pkg.DocGate):
        async def run(self, context):
            return GateResult(
                name="document", passed=False,
                failures=["document remains incomplete"],
            )

    monkeypatch.setattr(
        "research_assistant.pipeline.runner._run_stage_agent", fake_agent,
    )
    monkeypatch.setattr(gates_pkg, "DocGate", FailingDocGate)

    from research_assistant.pipeline.runner import run_pipeline

    out_dir = tmp_path / "run"
    stale_final = out_dir / "final" / "manuscript.docx"
    stale_final.parent.mkdir(parents=True)
    stale_final.write_bytes(b"stale")

    updates = await _collect(run_pipeline(
        "q", "m", tmp_path, out_dir, max_revision_rounds=1,
    ))

    result = next(u for u in updates if u.get("type") == "result")
    assert result["status"] == "failed"
    assert result["compilation_success"] is False
    assert "[document] document remains incomplete" in result["errors"]
    assert not stale_final.exists()
    assert (out_dir / "drafts" / "v1_draft.docx").exists()
    assert "review" not in calls

    session = SessionStore(out_dir)
    assert session.state.status == "failed"
    assert session.stage_status("gates") == "failed"


@pytest.mark.asyncio
async def test_pipeline_resume_skips_completed_stages(tmp_path, monkeypatch, gate_stub):
    calls: list = []
    monkeypatch.setattr(
        "research_assistant.pipeline.runner._run_stage_agent",
        _make_fake_agent(calls, {}),
    )
    from research_assistant.pipeline.runner import run_pipeline

    out_dir = tmp_path / "run"
    await _collect(run_pipeline("q", "m", tmp_path, out_dir))
    first_run_calls = list(calls)
    assert any(c.startswith("research_") for c in first_run_calls)

    # Simulate crash after assembly: wipe drafts + final but keep artifacts.
    import shutil
    shutil.rmtree(out_dir / "drafts")
    shutil.rmtree(out_dir / "final")

    calls.clear()
    await _collect(run_pipeline("q", "m", tmp_path, out_dir))

    # Research/figure/plan stages must NOT rerun — artifacts were valid.
    assert not any(c.startswith("research_") for c in calls)
    assert not any(c.startswith("figure_") for c in calls)
    assert "plan" not in calls
    assert "assemble" in calls  # its artifact was deleted -> must regenerate
    assert (out_dir / "final" / "manuscript.docx").exists()


@pytest.mark.asyncio
async def test_pipeline_cancel_before_start(tmp_path, monkeypatch, gate_stub):
    import asyncio
    cancel = asyncio.Event()
    cancel.set()
    calls: list = []
    monkeypatch.setattr(
        "research_assistant.pipeline.runner._run_stage_agent",
        _make_fake_agent(calls, {}),
    )
    from research_assistant.pipeline.runner import run_pipeline

    out_dir = tmp_path / "run"
    await _collect(run_pipeline(
        "q", "m", tmp_path, out_dir, cancel_event=cancel,
    ))
    assert calls == []
    session = SessionStore(out_dir)
    assert session.state.status == "cancelled"


@pytest.mark.asyncio
async def test_pipeline_records_budget_snapshot(tmp_path, monkeypatch, gate_stub):
    from research_assistant.kernel.budget import BudgetLimits
    calls: list = []
    monkeypatch.setattr(
        "research_assistant.pipeline.runner._run_stage_agent",
        _make_fake_agent(calls, {}),
    )
    from research_assistant.pipeline.runner import run_pipeline

    out_dir = tmp_path / "run"
    await _collect(run_pipeline(
        "q", "claude-sonnet-5", tmp_path, out_dir,
        budget_limits=BudgetLimits(max_turns=100),
    ))
    state = SessionStore(out_dir).state
    assert state.budget["model"] == "claude-sonnet-5"
    assert state.budget["limits"]["max_turns"] == 100


# ---------------------------------------------------------------------------
# R12 P1：执行契约注入（阶段子代理 system_prompt 收口处）
# ---------------------------------------------------------------------------

class TestStageContractInjection:
    async def test_stage_system_prompt_carries_frozen_contract(self, monkeypatch):
        import sys

        from research_assistant.pipeline import runner

        captured: dict = {}

        async def fake_run_agent(**kwargs):
            captured["system_prompt"] = kwargs["system_prompt"]
            return object()  # runner 原样返回，占位即可

        class _FakeClient:
            async def close(self):
                pass

        monkeypatch.setattr(runner, "run_agent", fake_run_agent)
        monkeypatch.setattr(runner, "create_llm_client", lambda **kw: _FakeClient())
        monkeypatch.setattr(sys, "frozen", True, raising=False)

        await runner._run_stage_agent(
            "plan",
            "do the plan",
            "BASE STAGE PROMPT",
            model="m",
            work_dir=Path("."),
            api_key=None,
            base_url=None,
            provider=None,
            budget=None,
            hooks=None,
            cancel_event=None,
        )
        # 原阶段提示词保留在前，契约追加在后
        assert captured["system_prompt"].startswith("BASE STAGE PROMPT")
        assert "run_script" in captured["system_prompt"]
        assert "sys.executable" in captured["system_prompt"]


# ---------------------------------------------------------------------------
# R12 P2/B1：SessionState.outputs_dir 字段（chat 模式产物目录落盘）
# ---------------------------------------------------------------------------

class TestSessionStateOutputsDir:
    def test_field_roundtrip(self, tmp_path):
        store = SessionStore.create(tmp_path / "run", "q", "m")
        store.state.outputs_dir = "outputs/20260823_100000_demo"
        store.save()
        reloaded = SessionStore(tmp_path / "run")
        assert reloaded.state.outputs_dir == "outputs/20260823_100000_demo"

    def test_old_run_json_without_field_loads(self, tmp_path):
        run_dir = tmp_path / "run"
        run_dir.mkdir()
        (run_dir / "run.json").write_text(json.dumps({
            "schema_version": 1, "session_id": "run", "query": "q",
            "model": "m", "mode": "chat", "status": "complete",
        }), encoding="utf-8")
        assert SessionStore(run_dir).state.outputs_dir == ""


# ---------------------------------------------------------------------------
# R12 P2/B5：任务侧写入归巢——write_anchor=paper 目录，exec_cwd 语义不变
# ---------------------------------------------------------------------------


class TestStageWriteAnchor:
    async def _run_stage(self, monkeypatch, tmp_path, **extra):
        from research_assistant.pipeline import runner

        captured: dict = {}

        async def fake_run_agent(**kwargs):
            captured["tools"] = kwargs["tools"]
            return object()

        class _FakeClient:
            async def close(self):
                pass

        monkeypatch.setattr(runner, "run_agent", fake_run_agent)
        monkeypatch.setattr(runner, "create_llm_client", lambda **kw: _FakeClient())

        await runner._run_stage_agent(
            "plan", "p", "sp", model="m",
            work_dir=tmp_path, api_key=None, base_url=None, provider=None,
            budget=None, hooks=None, cancel_event=None, **extra)
        return captured

    async def test_write_anchor_threads_paper_dir(self, monkeypatch, tmp_path):
        paper_dir = tmp_path / "writing_outputs" / "20260823_paper"
        tools = (await self._run_stage(
            monkeypatch, tmp_path, write_anchor=str(paper_dir)))["tools"]
        # 相对写入归巢论文目录；exec_cwd 不动（阶段提示词全绝对路径，
        # CWD=根 的既有语义保持）
        assert tools.write_anchor == str(paper_dir)
        assert tools.exec_cwd == str(tmp_path)

    async def test_default_keeps_legacy_no_anchor(self, monkeypatch, tmp_path):
        tools = (await self._run_stage(monkeypatch, tmp_path))["tools"]
        assert tools.write_anchor is None

    async def test_run_pipeline_anchors_all_stages_to_output_dir(
            self, tmp_path, monkeypatch, gate_stub):
        """端到端：流水线每个阶段的 write_anchor 都指向论文输出目录。"""
        calls: list = []
        anchors: list = []
        inner = _make_fake_agent(calls, {})

        async def spy(stage_name, prompt, system_prompt, *, write_anchor=None, **kw):
            anchors.append(write_anchor)
            kw.pop("write_anchor", None)
            return await inner(stage_name, prompt, system_prompt, **kw)

        monkeypatch.setattr(
            "research_assistant.pipeline.runner._run_stage_agent", spy)
        from research_assistant.pipeline.runner import run_pipeline

        out_dir = tmp_path / "writing_outputs" / "run1"
        await _collect(run_pipeline(
            query="write a fem paper", model="test-model",
            work_dir=tmp_path, output_dir=out_dir,
        ))

        assert calls  # 各阶段确实运行
        assert set(anchors) == {str(out_dir)}

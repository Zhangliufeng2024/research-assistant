"""Tests for research_assistant.utils."""

import tempfile
from pathlib import Path

from research_assistant.utils import (
    find_existing_papers,
    detect_paper_reference,
    scan_paper_directory,
    count_citations_in_bib,
    count_words_in_tex,
    extract_title_from_tex,
    extract_citation_style,
    count_words_in_docx,
    extract_title_from_docx,
)


class TestFindExistingPapers:
    def test_empty_folder(self, tmp_path):
        assert find_existing_papers(tmp_path) == []

    def test_nonexistent_folder(self, tmp_path):
        assert find_existing_papers(tmp_path / "nope") == []

    def test_returns_sorted_by_mtime(self, tmp_path):
        import time, os
        old = tmp_path / "20240101_000000_old"
        old.mkdir()
        # Force different mtime by backdating the old directory
        old_time = time.time() - 100
        os.utime(old, (old_time, old_time))
        new = tmp_path / "20240102_000000_new"
        new.mkdir()
        papers = find_existing_papers(tmp_path)
        assert len(papers) == 2
        assert "new" in papers[0]["name"]

    def test_ignores_files(self, tmp_path):
        (tmp_path / "not_a_dir.txt").write_text("hi")
        assert find_existing_papers(tmp_path) == []


class TestDetectPaperReference:
    def _make_papers(self, tmp_path, names):
        papers = []
        for name in names:
            d = tmp_path / name
            d.mkdir()
            papers.append({"path": d, "name": name, "mtime": d.stat().st_mtime})
        return papers

    def test_no_papers(self):
        assert detect_paper_reference("continue the paper", []) is None

    def test_new_paper_keyword_returns_none(self, tmp_path):
        papers = self._make_papers(tmp_path, ["20240101_000000_quantum_stuff"])
        assert detect_paper_reference("new paper on chemistry", papers) is None

    def test_continuation_keyword_returns_most_recent(self, tmp_path):
        papers = self._make_papers(tmp_path, ["20240101_000000_quantum_stuff"])
        result = detect_paper_reference("continue the paper", papers)
        assert result == papers[0]["path"]

    def test_topic_match(self, tmp_path):
        papers = self._make_papers(tmp_path, [
            "20240101_000000_quantum_computing",
            "20240102_000000_machine_learning",
        ])
        result = detect_paper_reference("find the machine learning paper", papers)
        assert "machine_learning" in str(result)

    def test_no_match_returns_none(self, tmp_path):
        papers = self._make_papers(tmp_path, ["20240101_000000_quantum_stuff"])
        assert detect_paper_reference("hello world", papers) is None


class TestScanPaperDirectory:
    def test_nonexistent_dir(self, tmp_path):
        result = scan_paper_directory(tmp_path / "nope")
        assert result["pdf_final"] is None
        assert result["figures"] == []

    def test_full_structure(self, tmp_path):
        (tmp_path / "final").mkdir()
        (tmp_path / "final" / "manuscript.pdf").write_bytes(b"pdf")
        (tmp_path / "final" / "manuscript.tex").write_text("tex")
        (tmp_path / "drafts").mkdir()
        (tmp_path / "drafts" / "v1_draft.tex").write_text("tex")
        (tmp_path / "drafts" / "v1_draft.pdf").write_bytes(b"pdf")
        (tmp_path / "references").mkdir()
        (tmp_path / "references" / "references.bib").write_text("@article{}")
        (tmp_path / "figures").mkdir()
        (tmp_path / "figures" / "fig1.png").write_bytes(b"png")
        (tmp_path / "data").mkdir()
        (tmp_path / "data" / "data.csv").write_text("a,b")
        (tmp_path / "sources").mkdir()
        (tmp_path / "sources" / "notes.md").write_text("notes")
        (tmp_path / "progress.md").write_text("progress")
        (tmp_path / "SUMMARY.md").write_text("summary")

        result = scan_paper_directory(tmp_path)
        assert result["pdf_final"] is not None
        assert result["tex_final"] is not None
        assert len(result["tex_drafts"]) == 1
        assert len(result["pdf_drafts"]) == 1
        assert result["bibliography"] is not None
        assert len(result["figures"]) == 1
        assert len(result["data"]) == 1
        assert len(result["sources"]) == 1
        assert result["progress_log"] is not None
        assert result["summary"] is not None


class TestCountCitationsInBib:
    def test_none_file(self):
        assert count_citations_in_bib(None) == 0

    def test_nonexistent_file(self):
        assert count_citations_in_bib("/nonexistent/path.bib") == 0

    def test_counts_entries(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text(
            "@article{smith2020,\n  title={A},\n}\n"
            "@inproceedings{jones2021,\n  title={B},\n}\n"
            "@comment{ignore this}\n"
        )
        assert count_citations_in_bib(str(bib)) == 2

    def test_ignores_comments(self, tmp_path):
        bib = tmp_path / "refs.bib"
        bib.write_text("@comment{skip}\n% @article{commented}\n@book{real,}\n")
        assert count_citations_in_bib(str(bib)) == 1


class TestCountWordsInTex:
    def test_none_file(self):
        assert count_words_in_tex(None) is None

    def test_simple_content(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text("Hello world this is a test")
        count = count_words_in_tex(str(tex))
        assert count == 6

    def test_strips_commands(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text(r"\textbf{bold} normal \emph{italic} text")
        count = count_words_in_tex(str(tex))
        assert count is not None and count >= 2

    def test_strips_comments(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text("real words here\n% this is a comment\nmore words")
        count = count_words_in_tex(str(tex))
        assert count is not None and count >= 4


class TestExtractTitleFromTex:
    def test_none_file(self):
        assert extract_title_from_tex(None) is None

    def test_extracts_title(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text(r"\title{My Amazing Paper}")
        assert extract_title_from_tex(str(tex)) == "My Amazing Paper"

    def test_multiline_title(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text("\\title{Line One\nLine Two}")
        title = extract_title_from_tex(str(tex))
        assert title is not None
        assert "Line One" in title

    def test_no_title(self, tmp_path):
        tex = tmp_path / "paper.tex"
        tex.write_text(r"\begin{document} no title here \end{document}")
        assert extract_title_from_tex(str(tex)) is None


class TestScanPaperDirectoryDocx:
    def test_docx_in_final(self, tmp_path):
        (tmp_path / "final").mkdir()
        (tmp_path / "final" / "manuscript.docx").write_bytes(b"pk")
        result = scan_paper_directory(tmp_path)
        assert result["docx_final"] is not None

    def test_docx_in_drafts(self, tmp_path):
        (tmp_path / "drafts").mkdir()
        (tmp_path / "drafts" / "v1_draft.docx").write_bytes(b"pk")
        result = scan_paper_directory(tmp_path)
        assert len(result["docx_drafts"]) == 1


class TestCountWordsInDocx:
    def test_none_file(self):
        assert count_words_in_docx(None) is None

    def test_nonexistent_file(self):
        assert count_words_in_docx("/nonexistent/file.docx") is None

    def test_valid_docx(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("hello world this is a test")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        count = count_words_in_docx(str(path))
        assert count == 6


class TestExtractTitleFromDocx:
    def test_none_file(self):
        assert extract_title_from_docx(None) is None

    def test_nonexistent_file(self):
        assert extract_title_from_docx("/nonexistent/file.docx") is None

    def test_title_style(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph("My Great Paper")
        p.style = doc.styles["Title"]
        doc.add_paragraph("Some content")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        assert extract_title_from_docx(str(path)) == "My Great Paper"

    def test_fallback_first_paragraph(self, tmp_path):
        from docx import Document
        doc = Document()
        doc.add_paragraph("First Line")
        doc.add_paragraph("Second Line")
        path = tmp_path / "test.docx"
        doc.save(str(path))
        assert extract_title_from_docx(str(path)) == "First Line"


class TestExtractCitationStyle:
    def test_default(self):
        assert extract_citation_style(None) == "BibTeX"
